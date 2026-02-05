"""Gemini 3 Pro Image による画像生成"""

from __future__ import annotations

import base64
import logging
import re
from dataclasses import dataclass, field
from pathlib import Path

from src.utils.config import get_env_var, load_settings
from src.utils.exceptions import ConfigurationError, ImageGenerationError
from src.utils.retry import with_retry

logger = logging.getLogger(__name__)

# リトライ設定
MAX_RETRIES = 3
BASE_DELAY = 2.0


@dataclass
class ImagePrompt:
    """画像プロンプトデータ"""

    number: int
    start_time: str
    end_time: str
    prompt: str


@dataclass
class ImagePromptList:
    """画像プロンプト一覧"""

    filename: str
    prompts: list[ImagePrompt] = field(default_factory=list)

    @property
    def total_images(self) -> int:
        return len(self.prompts)


class ImageGenerator:
    """Gemini 3 Pro Image による画像生成クライアント"""

    # プロンプトパターン: [1] 0:00-0:15 | プロンプト
    PROMPT_PATTERN = re.compile(
        r"\[(\d+)\]\s*(\d+:\d+)-(\d+:\d+)\s*\|\s*(.+)"
    )

    def __init__(self) -> None:
        self._client = None
        self._settings = load_settings()

    def _get_client(self):
        """クライアントを遅延初期化"""
        if self._client is None:
            api_key = get_env_var("GOOGLE_API_KEY")
            if not api_key:
                raise ConfigurationError(
                    "GOOGLE_API_KEY が設定されていません。"
                    ".envファイルまたは環境変数を確認してください。"
                )

            try:
                import google.genai as genai

                self._client = genai.Client(api_key=api_key)
                logger.info("Gemini クライアントを初期化しました")
            except Exception as e:
                raise ImageGenerationError(
                    f"Gemini クライアントの初期化に失敗: {e}",
                    original_error=e,
                )
        return self._client

    def parse_prompt_file(self, file_path: str | Path) -> ImagePromptList:
        """プロンプトファイルを解析

        Args:
            file_path: プロンプトファイルのパス

        Returns:
            パース済みのプロンプト一覧
        """
        file_path = Path(file_path)

        if file_path.suffix.lower() == ".docx":
            from docx import Document

            doc = Document(file_path)
            content = "\n".join(para.text for para in doc.paragraphs)
        else:
            with open(file_path, encoding="utf-8") as f:
                content = f.read()

        return self.parse_prompt_text(content, file_path.name)

    def parse_prompt_text(
        self, content: str, filename: str = "prompts.txt"
    ) -> ImagePromptList:
        """プロンプトテキストを解析

        Args:
            content: プロンプトテキスト
            filename: ファイル名（識別用）

        Returns:
            パース済みのプロンプト一覧

        対応形式:
        1. [番号] 開始時間-終了時間 | プロンプト
        2. 【画像生成プロンプト】の後にプロンプトが続く形式
        """
        prompt_list = ImagePromptList(filename=filename)

        # まず標準形式を試す
        for line in content.split("\n"):
            line = line.strip()
            if not line:
                continue

            match = self.PROMPT_PATTERN.match(line)
            if not match:
                continue

            prompt = ImagePrompt(
                number=int(match.group(1)),
                start_time=match.group(2),
                end_time=match.group(3),
                prompt=match.group(4).strip(),
            )
            prompt_list.prompts.append(prompt)

        # 標準形式で見つからなかった場合、【画像生成プロンプト】形式を試す
        if len(prompt_list.prompts) == 0:
            prompt_list = self._parse_slide_format(content, filename)

        return prompt_list

    def _parse_slide_format(self, content: str, filename: str) -> ImagePromptList:
        """【画像生成プロンプト】形式のファイルを解析

        形式例:
        1. タイトル
        【画像生成プロンプト】
        プロンプト内容...
        """
        prompt_list = ImagePromptList(filename=filename)
        lines = content.split("\n")

        current_number = 0
        in_prompt_section = False
        current_prompt_lines = []

        for line in lines:
            line_stripped = line.strip()

            # 番号付きセクションの検出 (1. xxx, 2. xxx, etc.)
            number_match = re.match(r'^(\d+)\.\s*', line_stripped)
            if number_match:
                # 前のプロンプトを保存
                if current_number > 0 and current_prompt_lines:
                    prompt_text = " ".join(current_prompt_lines).strip()
                    if prompt_text:
                        interval = 10  # 10秒/画像
                        start_sec = (current_number - 1) * interval
                        end_sec = current_number * interval
                        prompt = ImagePrompt(
                            number=current_number,
                            start_time=f"{start_sec // 60}:{start_sec % 60:02d}",
                            end_time=f"{end_sec // 60}:{end_sec % 60:02d}",
                            prompt=prompt_text,
                        )
                        prompt_list.prompts.append(prompt)

                current_number = int(number_match.group(1))
                in_prompt_section = False
                current_prompt_lines = []
                continue

            # 【画像生成プロンプト】セクションの開始
            if "【画像生成プロンプト】" in line_stripped:
                in_prompt_section = True
                current_prompt_lines = []
                continue

            # 他のセクションマーカーでプロンプトセクション終了
            if line_stripped.startswith("【") and "】" in line_stripped:
                in_prompt_section = False
                continue

            # プロンプトセクション内のテキストを収集
            if in_prompt_section and line_stripped:
                current_prompt_lines.append(line_stripped)

        # 最後のプロンプトを保存
        if current_number > 0 and current_prompt_lines:
            prompt_text = " ".join(current_prompt_lines).strip()
            if prompt_text:
                interval = 10
                start_sec = (current_number - 1) * interval
                end_sec = current_number * interval
                prompt = ImagePrompt(
                    number=current_number,
                    start_time=f"{start_sec // 60}:{start_sec % 60:02d}",
                    end_time=f"{end_sec // 60}:{end_sec % 60:02d}",
                    prompt=prompt_text,
                )
                prompt_list.prompts.append(prompt)

        return prompt_list

    def parse_uploaded_file(self, uploaded_file) -> ImagePromptList:
        """Streamlitのアップロードファイルを解析

        Args:
            uploaded_file: StreamlitのUploadedFileオブジェクト

        Returns:
            パース済みのプロンプト一覧
        """
        filename = uploaded_file.name

        if filename.lower().endswith(".docx"):
            from io import BytesIO

            from docx import Document

            doc = Document(BytesIO(uploaded_file.getvalue()))
            content = "\n".join(para.text for para in doc.paragraphs)
        else:
            content = uploaded_file.getvalue().decode("utf-8")

        return self.parse_prompt_text(content, filename)

    def generate(self, prompt: str, output_path: str | Path) -> Path:
        """プロンプトから画像を生成

        Args:
            prompt: 画像生成プロンプト
            output_path: 出力ファイルパス

        Returns:
            出力ファイルのパス

        Raises:
            ImageGenerationError: 画像生成に失敗した場合
            ConfigurationError: APIキーが設定されていない場合
        """
        return self._generate_with_retry(prompt, output_path)

    @with_retry(max_retries=MAX_RETRIES, base_delay=BASE_DELAY)
    def _generate_with_retry(self, prompt: str, output_path: str | Path) -> Path:
        """リトライ付き画像生成（内部メソッド）"""
        try:
            client = self._get_client()
            image_settings = self._settings.get("image_generation", {})
            model = image_settings.get("model", "gemini-3-pro-image-preview")

            logger.info("画像生成開始: model=%s, prompt_length=%d", model, len(prompt))

            # Gemini 3 Pro Image で画像生成
            from google.genai import types

            response = client.models.generate_content(
                model=model,
                contents=prompt,
                config=types.GenerateContentConfig(
                    image_config=types.ImageConfig(
                        aspect_ratio="16:9",
                    )
                ),
            )

            output_path = Path(output_path)
            output_path.parent.mkdir(parents=True, exist_ok=True)

            # レスポンスから画像データを取得
            image_parts = [part for part in response.parts if hasattr(part, 'inline_data') and part.inline_data]

            if image_parts:
                image = image_parts[0].as_image()
                image.save(str(output_path))
                logger.info("画像生成完了: %s", output_path)
                return output_path

            # テキストレスポンスのみの場合はログ出力
            for part in response.parts:
                if hasattr(part, 'text') and part.text:
                    logger.warning("テキストレスポンス受信: %s", part.text[:200])

            raise ImageGenerationError("レスポンスに画像データが含まれていません")

        except ConfigurationError:
            raise
        except ImageGenerationError:
            raise
        except Exception as e:
            error_msg = f"画像生成に失敗しました: {e}"
            logger.error(error_msg)
            raise ImageGenerationError(error_msg, original_error=e)
