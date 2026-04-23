"""台本パーサー

台本ファイル（Word/テキスト）を解析し、構造化データに変換する。
"""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document


@dataclass
class Line:
    """セリフデータ"""

    number: int
    speaker: str
    text: str
    original_text: str
    scene_description: str | None = None
    reading_hints: dict[str, str] = field(default_factory=dict)


@dataclass
class Script:
    """台本データ"""

    filename: str
    lines: list[Line] = field(default_factory=list)

    @property
    def total_lines(self) -> int:
        return len(self.lines)


class ScriptParser:
    """台本パーサー"""

    # 情景補足パターン: (ため息をついて)（驚いた表情で）など（半角・全角括弧対応）
    SCENE_PATTERN = re.compile(r"[（(]([^)）]+)[)）]")

    # 読み仮名パターン: {漢字|読み} など
    READING_PATTERN = re.compile(r"\{([^|]+)\|([^}]+)\}")

    # タイムスタンプ・章見出しパターン: 【2:30】 失敗のメカニズム解説① など
    TIMESTAMP_SECTION_PATTERN = re.compile(r"\s*【[^】]*】.*$")

    # 非セリフ行パターン（タイトル・見出し・装飾行）
    NON_DIALOGUE_PATTERNS = [
        re.compile(r"^【.*】"),                    # 【チャプター名】
        re.compile(r"^〈.*〉"),                    # 〈補足〉
        re.compile(r"^■|^●|^◆|^▶|^◎|^★|^☆"),     # 記号見出し
        re.compile(r"^━+|^─+|^＝+|^=+|^-{3,}"),   # 装飾線
        re.compile(r"^#{1,6}\s"),                  # Markdownの見出し
        re.compile(r"^タイトル"),                   # タイトル行（タイトル, タイトル:, タイトル案）
        re.compile(r"^動画タイトル"),               # 動画タイトル
        re.compile(r"^テーマ[：:]"),               # テーマ行
        re.compile(r"^台本[：:]"),                  # 台本ラベル
        re.compile(r"^\d+[：:]\d+[〜~～]\d+[：:]\d+"),  # タイムレンジ 0:00〜1:30
        re.compile(r"^※"),                         # 注釈行
        re.compile(r"^[\[（\(].*[）\)\]]$"),       # 全体が括弧の行（ト書き）
        re.compile(r"^画像生成プロンプト|^BGM|^SE[：:]"),  # メタ指示
        re.compile(r"^サムネ"),                     # サムネイル, サムネ案, サムネテキスト
        re.compile(r"^概要"),                       # 概要欄, 概要文
        re.compile(r"^ハッシュタグ"),               # ハッシュタグ
        re.compile(r"^動画説明"),                   # 動画説明
        re.compile(r"^(大|小|メイン)テキスト"),     # サムネ用テキストラベル
        re.compile(r"^テキスト\d*[：:]"),           # テキスト:, テキスト1:
        re.compile(r"^(テロップ|字幕|キャプション)"),  # テロップ・字幕指示
    ]

    # 話者パターン: speaker1:, Speaker 1:, ミオン：, アリイエ： など
    SPEAKER_PATTERN = re.compile(r"^(speaker\s*\d+):\s*(.+)$", re.IGNORECASE)
    # 話者のみのパターン（次の行にテキストがある場合）
    SPEAKER_ONLY_PATTERN = re.compile(r"^(speaker\s*\d+):\s*$", re.IGNORECASE)

    # キャラ名形式: 「名前：セリフ」または「名前: セリフ」（全角/半角コロン対応）
    # 名前は1〜10文字のひらがな・カタカナ・漢字・英字
    CHAR_NAME_PATTERN = re.compile(
        r"^([぀-ヿ㐀-䶵一-鿋豈-頻々〇〻\u3400-\u9FFFぁ-んァ-ヶa-zA-Zａ-ｚＡ-Ｚ]{1,10})[：:]\s*(.+)$"
    )
    CHAR_NAME_ONLY_PATTERN = re.compile(
        r"^([぀-ヿ㐀-䶵一-鿋豈-頻々〇〻\u3400-\u9FFFぁ-んァ-ヶa-zA-Zａ-ｚＡ-Ｚ]{1,10})[：:]\s*$"
    )

    # セクションラベル（これ以前をスキップし、以降をセリフとして解析）
    SCRIPT_START_PATTERNS = [
        re.compile(r"^本編シナリオ"),
        re.compile(r"^台本本文"),
        re.compile(r"^シナリオ本文"),
        re.compile(r"^本文"),
    ]

    # セクション区切り（これ以降のセリフ解析を停止）
    SCRIPT_END_PATTERNS = [
        re.compile(r"^画像生成プロンプト"),
        re.compile(r"^サムネ"),             # サムネイル, サムネ案, サムネテキスト
        re.compile(r"^エンディング"),
        re.compile(r"^概要欄"),
        re.compile(r"^ハッシュタグ"),
    ]

    # セクションヘッダー（この行＋後続のコンテンツ行をスキップ）
    SECTION_HEADER_PATTERNS = [
        re.compile(r"^タイトル\s*$"),          # 「タイトル」単体行
        re.compile(r"^タイトル[：:]?\s*$"),    # 「タイトル:」「タイトル：」
        re.compile(r"^動画タイトル\s*$"),      # 「動画タイトル」単体行
        re.compile(r"^サムネ"),                # サムネイル系セクション
    ]

    def parse_file(self, file_path: str | Path) -> Script:
        """ファイルを解析して台本データを返す

        Args:
            file_path: 台本ファイルのパス

        Returns:
            パース済みの台本データ
        """
        file_path = Path(file_path)

        if file_path.suffix.lower() == ".docx":
            content = self._read_docx(file_path)
        else:
            content = self._read_text(file_path)

        return self._parse_content(content, file_path.name)

    def parse_text(self, content: str, filename: str = "input.txt") -> Script:
        """テキストを解析して台本データを返す

        Args:
            content: 台本テキスト
            filename: ファイル名（識別用）

        Returns:
            パース済みの台本データ
        """
        return self._parse_content(content, filename)

    def parse_uploaded_file(self, uploaded_file) -> Script:
        """Streamlitのアップロードファイルを解析

        Args:
            uploaded_file: StreamlitのUploadedFileオブジェクト

        Returns:
            パース済みの台本データ
        """
        filename = uploaded_file.name

        if filename.lower().endswith(".docx"):
            from io import BytesIO

            doc = Document(BytesIO(uploaded_file.getvalue()))
            content = "\n".join(para.text for para in doc.paragraphs)
        else:
            content = uploaded_file.getvalue().decode("utf-8")

        return self._parse_content(content, filename)

    def _read_docx(self, file_path: Path) -> str:
        """Wordファイルを読み込む"""
        doc = Document(file_path)
        return "\n".join(para.text for para in doc.paragraphs)

    def _read_text(self, file_path: Path) -> str:
        """テキストファイルを読み込む"""
        with open(file_path, encoding="utf-8") as f:
            return f.read()

    def _is_non_dialogue(self, text: str) -> bool:
        """非セリフ行（タイトル・見出し・装飾）かどうか判定"""
        for pattern in self.NON_DIALOGUE_PATTERNS:
            if pattern.search(text):
                return True
        return False

    def _find_script_start(self, lines: list[str]) -> int:
        """セリフ開始位置を検出（「本編シナリオ」等のラベル以降）"""
        for i, line in enumerate(lines):
            stripped = line.strip()
            for pattern in self.SCRIPT_START_PATTERNS:
                if pattern.search(stripped):
                    return i + 1  # ラベルの次の行から
        return 0  # 見つからない場合は先頭から

    def _is_script_end(self, text: str) -> bool:
        """セリフ終了セクションかどうか判定"""
        for pattern in self.SCRIPT_END_PATTERNS:
            if pattern.search(text):
                return True
        return False

    def _is_section_header(self, text: str) -> bool:
        """セクションヘッダー（後続行もスキップすべき）かどうか判定"""
        for pattern in self.SECTION_HEADER_PATTERNS:
            if pattern.search(text):
                return True
        return False

    def _match_speaker(self, raw_line: str) -> tuple[str, str] | None:
        """話者+テキスト行にマッチ（speaker形式 + キャラ名形式）"""
        # speaker1: テキスト 形式
        match = self.SPEAKER_PATTERN.match(raw_line)
        if match:
            speaker = match.group(1).lower().replace(" ", "")
            return speaker, match.group(2)

        # キャラ名：テキスト 形式
        match = self.CHAR_NAME_PATTERN.match(raw_line)
        if match:
            return match.group(1), match.group(2)

        return None

    def _match_speaker_only(self, raw_line: str) -> str | None:
        """話者のみの行にマッチ"""
        match = self.SPEAKER_ONLY_PATTERN.match(raw_line)
        if match:
            return match.group(1).lower().replace(" ", "")

        match = self.CHAR_NAME_ONLY_PATTERN.match(raw_line)
        if match:
            return match.group(1)

        return None

    def _normalize_speaker(self, name: str, speaker_map: dict[str, str]) -> str:
        """キャラ名をspeaker1/speaker2に正規化"""
        if name.startswith("speaker"):
            return name
        if name not in speaker_map:
            # 新しいキャラ名 → 登場順にspeaker1, speaker2, ...
            idx = len(speaker_map) + 1
            speaker_map[name] = f"speaker{idx}"
        return speaker_map[name]

    def _parse_content(self, content: str, filename: str) -> Script:
        """コンテンツを解析（speaker形式 + キャラ名形式に対応）"""
        script = Script(filename=filename)
        line_number = 0
        lines = content.split("\n")
        speaker_map: dict[str, str] = {}  # キャラ名 → speaker1/speaker2

        # セリフ開始位置を検出
        start_idx = self._find_script_start(lines)
        i = start_idx

        while i < len(lines):
            raw_line = lines[i].strip()
            i += 1

            if not raw_line:
                continue

            # セリフ終了セクション
            if self._is_script_end(raw_line):
                break

            # 非セリフ行をスキップ
            if self._is_non_dialogue(raw_line):
                # セクションヘッダーの場合、後続のコンテンツ行もスキップ
                # 空行が現れるまで（最大2行）スキップする
                if self._is_section_header(raw_line):
                    skipped = 0
                    while i < len(lines) and skipped < 2:
                        peek = lines[i].strip()
                        if not peek:
                            if skipped > 0:
                                break  # コンテンツ後の空行＝セクション終了
                            i += 1
                            continue
                        if self._is_script_end(peek):
                            break
                        if self._is_non_dialogue(peek):
                            break
                        # コンテンツ行をスキップ（タイトルテキスト等）
                        i += 1
                        skipped += 1
                continue

            # 話者+テキスト行
            result = self._match_speaker(raw_line)
            if result:
                speaker_name, text = result
                speaker = self._normalize_speaker(speaker_name, speaker_map)
                # 後続の段落を次の話者まで集める（複数段落セリフに対応）
                text_parts = [text] if text else []
                while i < len(lines):
                    next_line = lines[i].strip()
                    if self._match_speaker(next_line) or self._match_speaker_only(next_line):
                        break
                    if self._is_script_end(next_line):
                        break
                    if next_line:
                        text_parts.append(next_line)
                    i += 1
                text = " ".join(text_parts).strip()
                if not text:
                    continue
                line_number += 1
            else:
                # 話者のみの行（次の行にテキスト）
                speaker_name = self._match_speaker_only(raw_line)
                if speaker_name:
                    speaker = self._normalize_speaker(speaker_name, speaker_map)
                    text_lines = []
                    while i < len(lines):
                        next_line = lines[i].strip()
                        if self._match_speaker(next_line) or self._match_speaker_only(next_line):
                            break
                        if self._is_script_end(next_line):
                            break
                        if next_line:
                            text_lines.append(next_line)
                        i += 1
                    text = " ".join(text_lines)
                    if not text:
                        continue
                    line_number += 1
                else:
                    continue

            # 情景補足を抽出・除去
            scene_match = self.SCENE_PATTERN.search(text)
            scene_description = scene_match.group(1) if scene_match else None
            clean_text = self.SCENE_PATTERN.sub("", text).strip()

            # 読み仮名を抽出
            reading_hints: dict[str, str] = {}
            for kanji, reading in self.READING_PATTERN.findall(clean_text):
                reading_hints[kanji] = reading

            # 読み仮名を展開（{漢字|読み} → 読み）
            final_text = self.READING_PATTERN.sub(r"\2", clean_text)

            # タイムスタンプ・章見出しを除去（TTS用テキストのみ）
            final_text = self.TIMESTAMP_SECTION_PATTERN.sub("", final_text).strip()

            if not final_text:
                line_number -= 1
                continue

            line = Line(
                number=line_number,
                speaker=speaker,
                text=final_text,
                original_text=text,
                scene_description=scene_description,
                reading_hints=reading_hints,
            )
            script.lines.append(line)

        # Speaker/キャラ名形式が見つからなかった場合、フォールバック
        if not script.lines:
            script = self._parse_content_fallback(content, filename)

        return script

    def _parse_content_fallback(self, content: str, filename: str) -> Script:
        """フォールバックパーサー: 番号付き行または通常のテキスト行を解析"""
        script = Script(filename=filename)
        lines = content.split("\n")

        # セリフ開始位置を検出
        start_idx = self._find_script_start(lines)

        # 番号付き行のパターン（1. テキスト、1: テキスト、1) テキスト など）
        numbered_pattern = re.compile(r"^(\d+)[.:\)）、\s]+(.+)$")

        line_number = 0
        current_speaker = "speaker1"  # デフォルト話者

        i = start_idx
        while i < len(lines):
            raw_line = lines[i].strip()
            i += 1
            if not raw_line:
                continue

            # セリフ終了セクション
            if self._is_script_end(raw_line):
                break

            # 非セリフ行をスキップ
            if self._is_non_dialogue(raw_line):
                # セクションヘッダーの場合、後続のコンテンツ行もスキップ
                # 空行が現れるまで（最大2行）スキップする
                if self._is_section_header(raw_line):
                    skipped = 0
                    while i < len(lines) and skipped < 2:
                        peek = lines[i].strip()
                        if not peek:
                            if skipped > 0:
                                break  # コンテンツ後の空行＝セクション終了
                            i += 1
                            continue
                        if self._is_script_end(peek):
                            break
                        if self._is_non_dialogue(peek):
                            break
                        # コンテンツ行をスキップ（タイトルテキスト等）
                        i += 1
                        skipped += 1
                continue

            # 番号付き行をチェック
            match = numbered_pattern.match(raw_line)
            if match:
                text = match.group(2).strip()
                # 番号付き行でも非セリフならスキップ
                if self._is_non_dialogue(text):
                    continue
                line_number += 1
                # 話者を交互に切り替え
                current_speaker = "speaker1" if line_number % 2 == 1 else "speaker2"
            else:
                # 通常のテキスト行（10文字以上かつ句読点を含む行のみ＝セリフらしい行）
                has_punctuation = any(c in raw_line for c in "。、！？!?」』)")
                if len(raw_line) >= 10 and has_punctuation:
                    line_number += 1
                    text = raw_line
                    current_speaker = "speaker1" if line_number % 2 == 1 else "speaker2"
                else:
                    continue

            # 情景補足を抽出・除去
            scene_match = self.SCENE_PATTERN.search(text)
            scene_description = scene_match.group(1) if scene_match else None
            clean_text = self.SCENE_PATTERN.sub("", text).strip()

            if not clean_text:
                line_number -= 1  # 空になった場合は行番号を戻す
                continue

            # 読み仮名を抽出
            reading_hints: dict[str, str] = {}
            for kanji, reading in self.READING_PATTERN.findall(clean_text):
                reading_hints[kanji] = reading

            # 読み仮名を展開
            final_text = self.READING_PATTERN.sub(r"\2", clean_text)

            # タイムスタンプ・章見出しを除去（TTS用テキストのみ）
            final_text = self.TIMESTAMP_SECTION_PATTERN.sub("", final_text).strip()

            if not final_text:
                line_number -= 1
                continue

            line = Line(
                number=line_number,
                speaker=current_speaker,
                text=final_text,
                original_text=text,
                scene_description=scene_description,
                reading_hints=reading_hints,
            )
            script.lines.append(line)

        return script
