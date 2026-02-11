"""動画生成エージェント - Streamlit メインアプリケーション"""

from __future__ import annotations

import base64
import json
import os
import shutil
import traceback
import zipfile
from datetime import datetime
from io import BytesIO
from pathlib import Path

import streamlit as st

from src.audio.tts import TTSClient
from src.bgm.beatoven import BeatovenClient
from src.image.generator import ImageGenerator
from src.parser.script import ScriptParser
from src.utils.config import get_env_var, get_gcp_credentials, load_settings, save_settings
from src.video.editor import Timeline, TimelineEntry, VideoEditor
from src.video.stock import StockVideoClient

# ページ設定
st.set_page_config(
    page_title="動画生成エージェント",
    page_icon="🎬",
    layout="wide",
    initial_sidebar_state="expanded",
)


def get_persistent_avatar_dir() -> Path:
    """永続的なアバター保存ディレクトリを取得"""
    # ドキュメントフォルダ内に保存（アプリ再起動後も保持される）
    home = os.path.expanduser("~")
    avatar_dir = Path(home) / "Documents" / "video-generator-output" / "avatars"
    avatar_dir.mkdir(parents=True, exist_ok=True)
    return avatar_dir


def restore_avatars_from_settings() -> None:
    """設定ファイルからアバター画像を復元（起動時に実行）"""
    settings = load_settings()
    avatar_dir = get_persistent_avatar_dir()

    for speaker_key in ["speaker1", "speaker2"]:
        speaker_settings = settings.get("speakers", {}).get(speaker_key, {})
        avatar_base64 = speaker_settings.get("avatar_base64")
        avatar_ext = speaker_settings.get("avatar_ext", "png")

        if avatar_base64:
            try:
                # Base64からデコード
                image_data = base64.b64decode(avatar_base64)
                avatar_path = avatar_dir / f"{speaker_key}.{avatar_ext}"

                # ファイルを復元（常に上書き - Base64が真のソース）
                with open(avatar_path, "wb") as f:
                    f.write(image_data)

                # パスも更新
                if "speakers" not in settings:
                    settings["speakers"] = {}
                if speaker_key not in settings["speakers"]:
                    settings["speakers"][speaker_key] = {}
                settings["speakers"][speaker_key]["avatar_path"] = str(avatar_path)
                save_settings(settings)

            except Exception as e:
                print(f"アバター復元エラー ({speaker_key}): {e}")


def save_avatar_to_settings(speaker_key: str, image_data: bytes, ext: str) -> None:
    """アバター画像をBase64で設定に保存"""
    settings = load_settings()

    if "speakers" not in settings:
        settings["speakers"] = {}
    if speaker_key not in settings["speakers"]:
        settings["speakers"][speaker_key] = {}

    # Base64エンコード
    avatar_base64 = base64.b64encode(image_data).decode("utf-8")
    settings["speakers"][speaker_key]["avatar_base64"] = avatar_base64
    settings["speakers"][speaker_key]["avatar_ext"] = ext

    # ファイルパスも保存
    avatar_dir = get_persistent_avatar_dir()
    avatar_path = avatar_dir / f"{speaker_key}.{ext}"
    settings["speakers"][speaker_key]["avatar_path"] = str(avatar_path)

    save_settings(settings)


# 起動時にアバターを復元
restore_avatars_from_settings()


def time_to_seconds(time_str: str) -> float:
    """時間文字列を秒に変換 (例: "1:30" -> 90.0)"""
    parts = time_str.split(":")
    if len(parts) == 2:
        return int(parts[0]) * 60 + int(parts[1])
    elif len(parts) == 3:
        return int(parts[0]) * 3600 + int(parts[1]) * 60 + int(parts[2])
    return 0.0


def count_script_items_from_content(content: str) -> int:
    """テキストから項数を検出（1, 2, 3... の番号から最大値を取得）"""
    import re
    max_item = 0

    # 各行をスキャン
    for line in content.split('\n'):
        line = line.strip()
        if not line:
            continue

        # 行頭の番号を検出（例: "1.", "1:", "1 ", "1）", "1)"）
        match = re.match(r'^(\d+)[.:\s）\)、]', line)
        if match:
            num = int(match.group(1))
            max_item = max(max_item, num)

    return max_item


def count_script_items(script) -> int:
    """台本から項数を検出（後方互換用）"""
    max_item = 0

    for line in script.lines:
        # 元のテキストから番号を検出
        text = line.original_text if hasattr(line, 'original_text') else line.text

        import re
        match = re.match(r'^(\d+)[.:\s）\)、]', text)
        if match:
            num = int(match.group(1))
            max_item = max(max_item, num)

    # 番号が見つからない場合は行数を返す
    return max_item if max_item > 0 else script.total_lines


def generate_image_prompts_from_script(script, num_images: int):
    """台本から画像プロンプトを自動生成"""
    from src.image.generator import ImagePrompt, ImagePromptList
    from src.utils.config import get_env_var
    import streamlit as st

    api_key = get_env_var("GOOGLE_API_KEY")
    use_ai_generation = bool(api_key)

    if not api_key:
        st.warning("⚠️ GOOGLE_API_KEY が未設定のため、簡易プロンプトを使用します")

    # ゼロ除算防止: 入力値の検証
    if num_images <= 0:
        num_images = max(1, script.total_lines if script.total_lines > 0 else 1)

    # 台本が空の場合の対応
    if not script.lines or len(script.lines) == 0:
        raise ValueError("台本が空です。セリフが含まれるファイルをアップロードしてください。")

    # 1セリフあたりの推定秒数（音声生成前なので概算）
    estimated_seconds_per_line = 5
    total_lines = max(1, script.total_lines)  # 0除算防止
    total_duration = total_lines * estimated_seconds_per_line

    # APIキーがある場合のみAI生成を試行
    if use_ai_generation:
        try:
            import google.genai as genai

            client = genai.Client(api_key=api_key)

            # 台本の全テキストを結合
            script_text = "\n".join([
                f"{line.number}. [{line.speaker}]: {line.text}"
                for line in script.lines
            ])

            prompt = f"""以下の台本を分析して、{num_images}枚の画像生成プロンプトを作成してください。

【台本】
{script_text}

【要件】
1. 各画像は台本の流れに沿ったシーンを表現する
2. プロンプトは日本語で、詳細な視覚的描写を含める
3. アニメ/イラスト風のスタイルを指定
4. 以下の形式で出力（各行1つのプロンプト）:

[番号] 開始時間-終了時間 | 日本語プロンプト

例:
[1] 0:00-0:10 | アニメ風、明るいスタジオで並んで座る2人のプロのニュースキャスター、フレンドリーな表情
[2] 0:10-0:20 | アニメ風、驚いた表情の女性キャラクターのクローズアップ、目を大きく見開いている

【注意】
- 時間は0:00から始め、{total_duration}秒程度で終わるように均等に配分
- 番号は1から{num_images}まで
- 各プロンプトは具体的で視覚的な描写を含める
- 台本の内容に合った適切なシーンを描写する
"""

            response = client.models.generate_content(
                model="gemini-2.0-flash",
                contents=prompt,
            )

            # レスポンスをパース
            from src.image.generator import ImageGenerator
            generator = ImageGenerator()
            result_text = response.text

            prompt_list = generator.parse_prompt_text(result_text, "auto_generated")

            if prompt_list.total_images > 0:
                return prompt_list

            # パースに失敗した場合はフォールバックへ
            st.warning(f"⚠️ AIレスポンスのパースに失敗。フォールバックを使用します。")

        except Exception as e:
            st.warning(f"⚠️ AI生成エラー: {e}。フォールバックを使用します。")

    # フォールバック: 台本から直接プロンプトを構築
    # ゼロ除算を防止
    if num_images <= 0:
        num_images = max(1, len(script.lines) if script.lines else 1)
    if total_duration <= 0:
        total_duration = max(num_images * 5, 5)  # デフォルト5秒/画像、最低5秒

    interval = max(1, total_duration // num_images)
    prompts = []

    # 台本が空の場合のフォールバック
    script_lines = script.lines if script.lines else []
    num_script_lines = len(script_lines)

    if num_script_lines == 0:
        # 台本が空の場合、デフォルトプロンプトを生成
        for i in range(num_images):
            start_sec = i * interval
            end_sec = (i + 1) * interval
            start_time = f"{start_sec // 60}:{start_sec % 60:02d}"
            end_time = f"{end_sec // 60}:{end_sec % 60:02d}"
            prompt_text = "アニメ風イラスト、カラフル、高品質、シーン背景"
            prompts.append(ImagePrompt(
                number=i + 1,
                start_time=start_time,
                end_time=end_time,
                prompt=prompt_text,
            ))
        return ImagePromptList(filename="auto_generated", prompts=prompts)

    # 各セリフからキーワードを抽出してプロンプトを生成
    lines_per_image = max(1, num_script_lines // num_images)

    for i in range(num_images):
        start_sec = i * interval
        end_sec = (i + 1) * interval
        start_time = f"{start_sec // 60}:{start_sec % 60:02d}"
        end_time = f"{end_sec // 60}:{end_sec % 60:02d}"

        # 対応するセリフからコンテキストを取得
        line_idx = min(i * lines_per_image, num_script_lines - 1)
        context = script_lines[line_idx].text[:100] if line_idx >= 0 else "シーン"

        # 日本語プロンプトを生成
        prompt_text = f"アニメ風イラスト、カラフル、高品質、シーン: {context}"

        prompts.append(ImagePrompt(
            number=i + 1,
            start_time=start_time,
            end_time=end_time,
            prompt=prompt_text,
        ))

    return ImagePromptList(filename="auto_generated", prompts=prompts)


def get_default_output_folder() -> str:
    """OSに応じたデフォルト出力フォルダを取得"""
    import platform

    # Docker環境かどうかを検出
    is_docker = os.path.exists("/.dockerenv") or os.environ.get("DOCKER_CONTAINER", False)

    if is_docker:
        # Docker環境では /app/output を使用（ホストにマウントされている前提）
        return "/app/output"

    system = platform.system()
    home = os.path.expanduser("~")

    if system == "Windows":
        # Windows: ドキュメントフォルダ内に作成
        docs_folder = os.path.join(home, "Documents", "video-generator-output")
        return docs_folder
    elif system == "Darwin":
        # macOS: ドキュメントフォルダ内に作成
        docs_folder = os.path.join(home, "Documents", "video-generator-output")
        return docs_folder
    else:
        # Linux: ホームディレクトリ内に作成
        return os.path.join(home, "video-generator-output")


def get_output_dir() -> Path:
    """出力ディレクトリを取得"""
    # セッション状態からカスタム出力先を取得（あれば）
    if "custom_output_folder" in st.session_state and st.session_state.custom_output_folder:
        output_folder = st.session_state.custom_output_folder
    else:
        settings = load_settings()
        configured_folder = settings.get("defaults", {}).get("output_folder", "")

        # 設定が "output"（相対パス）または空の場合は、OS別デフォルトを使用
        if not configured_folder or configured_folder == "output":
            output_folder = get_default_output_folder()
        else:
            output_folder = configured_folder

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    output_dir = Path(output_folder) / timestamp
    output_dir.mkdir(parents=True, exist_ok=True)
    return output_dir


def get_existing_output_folders() -> list[tuple[str, str]]:
    """既存の出力フォルダ一覧を取得（履歴からも取得）

    Returns:
        list of (folder_name, full_path) tuples
    """
    folders = []
    seen_names = set()

    # 1. 履歴から出力フォルダを取得（最優先）
    history = load_generation_history()
    for entry in history:
        output_dir = entry.get("output_dir", "")
        if output_dir:
            output_path = Path(output_dir)
            if output_path.exists():
                # audio, images, bgmのいずれかにファイルが存在するかチェック
                audio_dir = output_path / "audio"
                image_dir = output_path / "images"
                bgm_dir = output_path / "bgm"

                has_audio = audio_dir.exists() and any(audio_dir.glob("*.wav")) or any(audio_dir.glob("*.mp3")) if audio_dir.exists() else False
                has_images = image_dir.exists() and (any(image_dir.glob("*.png")) or any(image_dir.glob("*.jpg"))) if image_dir.exists() else False
                has_bgm = bgm_dir.exists() and (any(bgm_dir.glob("*.mp3")) or any(bgm_dir.glob("*.wav"))) if bgm_dir.exists() else False

                if has_audio or has_images or has_bgm:
                    folder_name = output_path.name
                    if folder_name not in seen_names:
                        folders.append((folder_name, str(output_path)))
                        seen_names.add(folder_name)

    # 2. 設定の出力フォルダからも取得
    if "custom_output_folder" in st.session_state and st.session_state.custom_output_folder:
        output_folder = st.session_state.custom_output_folder
    else:
        settings = load_settings()
        configured_folder = settings.get("defaults", {}).get("output_folder", "")
        # 設定が "output"（相対パス）または空の場合は、OS別デフォルトを使用
        if not configured_folder or configured_folder == "output":
            output_folder = get_default_output_folder()
        else:
            output_folder = configured_folder

    output_path = Path(output_folder)

    if output_path.exists():
        for folder in sorted(output_path.iterdir(), reverse=True):
            if folder.is_dir() and not folder.name.startswith("."):
                # audio, images, bgmのいずれかにファイルが存在するかチェック
                audio_dir = folder / "audio"
                image_dir = folder / "images"
                bgm_dir = folder / "bgm"

                has_audio = audio_dir.exists() and (any(audio_dir.glob("*.wav")) or any(audio_dir.glob("*.mp3"))) if audio_dir.exists() else False
                has_images = image_dir.exists() and (any(image_dir.glob("*.png")) or any(image_dir.glob("*.jpg"))) if image_dir.exists() else False
                has_bgm = bgm_dir.exists() and (any(bgm_dir.glob("*.mp3")) or any(bgm_dir.glob("*.wav"))) if bgm_dir.exists() else False

                if has_audio or has_images or has_bgm:
                    folder_name = folder.name
                    if folder_name not in seen_names:
                        folders.append((folder_name, str(folder)))
                        seen_names.add(folder_name)

    return folders


def load_existing_materials(folder_path_or_name: str) -> dict:
    """指定フォルダから素材を読み込む

    Args:
        folder_path_or_name: フルパスまたはフォルダ名
    """
    # フルパスかどうかを判定
    if os.path.isabs(folder_path_or_name) or folder_path_or_name.startswith("/"):
        folder_path = Path(folder_path_or_name)
    else:
        # フォルダ名の場合は設定から親フォルダを取得
        settings = load_settings()
        output_folder = settings.get("defaults", {}).get("output_folder", "output")
        folder_path = Path(output_folder) / folder_path_or_name

    result = {
        "audio_files": {},
        "images": {},
        "bgm": None,
        "videos": {},
    }

    # 音声ファイルを読み込み
    audio_dir = folder_path / "audio"
    if audio_dir.exists():
        for audio_file in audio_dir.glob("*.wav"):
            if audio_file.name == "full_audio.wav":
                result["audio_files"]["full"] = str(audio_file)
            else:
                # 001_speaker1.wav 形式から番号を抽出
                try:
                    num = int(audio_file.stem.split("_")[0])
                    result["audio_files"][num] = str(audio_file)
                except (ValueError, IndexError):
                    pass
        # MP3も対応
        for audio_file in audio_dir.glob("*.mp3"):
            if audio_file.name == "full_audio.mp3":
                result["audio_files"]["full"] = str(audio_file)

    # 画像ファイルを読み込み
    image_dir = folder_path / "images"
    if image_dir.exists():
        for image_file in image_dir.glob("*.png"):
            try:
                num = int(image_file.stem.split("_")[0])
                result["images"][num] = str(image_file)
            except (ValueError, IndexError):
                pass
        for image_file in image_dir.glob("*.jpg"):
            try:
                num = int(image_file.stem.split("_")[0])
                result["images"][num] = str(image_file)
            except (ValueError, IndexError):
                pass

    # BGMファイルを読み込み
    bgm_dir = folder_path / "bgm"
    if bgm_dir.exists():
        for bgm_file in bgm_dir.glob("*.mp3"):
            result["bgm"] = str(bgm_file)
            break
        if not result["bgm"]:
            for bgm_file in bgm_dir.glob("*.wav"):
                result["bgm"] = str(bgm_file)
                break

    # 背景動画ファイルを読み込み
    video_dir = folder_path / "videos" / "backgrounds"
    if video_dir.exists():
        for video_file in video_dir.glob("*.mp4"):
            try:
                num = int(video_file.stem.split("_")[0])
                result["videos"][num] = str(video_file)
            except (ValueError, IndexError):
                pass

    return result


def get_history_file_path() -> Path:
    """履歴ファイルのパスを取得"""
    # セッション状態のカスタム出力先を優先
    if "custom_output_folder" in st.session_state and st.session_state.custom_output_folder:
        output_folder = st.session_state.custom_output_folder
    else:
        settings = load_settings()
        configured_folder = settings.get("defaults", {}).get("output_folder", "")
        # 設定が "output"（相対パス）または空の場合は、OS別デフォルトを使用
        if not configured_folder or configured_folder == "output":
            output_folder = get_default_output_folder()
        else:
            output_folder = configured_folder

    history_path = Path(output_folder) / "generation_history.json"

    # 親ディレクトリが存在しない場合は作成
    history_path.parent.mkdir(parents=True, exist_ok=True)

    return history_path


def load_generation_history() -> list[dict]:
    """生成履歴を読み込む"""
    history_file = get_history_file_path()
    if history_file.exists():
        try:
            with open(history_file, encoding="utf-8") as f:
                return json.load(f)
        except (json.JSONDecodeError, IOError):
            return []
    return []


def save_generation_history(history: list[dict]) -> None:
    """生成履歴を保存"""
    history_file = get_history_file_path()
    history_file.parent.mkdir(parents=True, exist_ok=True)
    with open(history_file, "w", encoding="utf-8") as f:
        json.dump(history, f, ensure_ascii=False, indent=2)


def log_error_to_file(output_dir: Path | str, error_type: str, error_message: str, details: str = "") -> None:
    """エラーをファイルに記録（例外を握りつぶして続行を妨げない）"""
    try:
        output_dir = Path(output_dir) if isinstance(output_dir, str) else output_dir
        output_dir.mkdir(parents=True, exist_ok=True)
        error_log_path = output_dir / "error_log.txt"
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

        with open(error_log_path, "a", encoding="utf-8") as f:
            f.write(f"{'='*60}\n")
            f.write(f"[{timestamp}] {error_type}\n")
            f.write(f"{'='*60}\n")
            f.write(f"エラーメッセージ: {error_message}\n")
            if details:
                f.write(f"\n詳細:\n{details}\n")
            f.write("\n")
    except Exception as log_err:
        # ログ書き込みエラーは無視して続行
        print(f"エラーログ書き込み失敗: {log_err}")


def read_error_log(output_dir: Path) -> str | None:
    """エラーログを読み込む"""
    error_log_path = output_dir / "error_log.txt"
    if error_log_path.exists():
        with open(error_log_path, encoding="utf-8") as f:
            return f.read()
    return None


def create_history_entry(output_dir: str, status: str = "in_progress") -> dict:
    """履歴エントリを作成"""
    return {
        "id": datetime.now().strftime("%Y%m%d_%H%M%S"),
        "output_dir": output_dir,
        "status": status,  # "in_progress", "completed", "interrupted"
        "created_at": datetime.now().isoformat(),
        "updated_at": datetime.now().isoformat(),
        "progress": {
            "script_parsed": False,
            "audio_generated": False,
            "images_generated": False,
            "bgm_generated": False,
            "video_generated": False,
        },
        "files": {
            "script": None,
            "script_file": None,  # 台本ファイルパス
            "prompts": None,
            "prompts_file": None,  # プロンプトファイルパス
            "audio_files": {},
            "images": {},
            "bgm": None,
            "videos": [],
        },
        "settings": {
            "output_mode": None,
            "output_formats": [],
        },
    }


def save_script_to_output(script, output_dir: Path) -> Path | None:
    """台本を出力フォルダに保存"""
    try:
        script_path = output_dir / "script_backup.json"
        script_data = {
            "filename": script.filename,
            "lines": [
                {
                    "number": line.number,
                    "speaker": line.speaker,
                    "text": line.text,
                    "scene_description": line.scene_description,
                    "original_text": getattr(line, "original_text", line.text),
                }
                for line in script.lines
            ],
            "total_lines": script.total_lines,
        }
        with open(script_path, "w", encoding="utf-8") as f:
            json.dump(script_data, f, ensure_ascii=False, indent=2)
        return script_path
    except Exception as e:
        print(f"台本保存エラー: {e}")
        return None


def load_script_from_output(output_dir: Path):
    """出力フォルダから台本を読み込み"""
    from src.parser.script import Script, Line

    script_path = output_dir / "script_backup.json"
    if not script_path.exists():
        return None

    try:
        with open(script_path, encoding="utf-8") as f:
            data = json.load(f)

        lines = []
        for line_data in data.get("lines", []):
            line = Line(
                number=line_data["number"],
                speaker=line_data["speaker"],
                text=line_data["text"],
                original_text=line_data.get("original_text", line_data["text"]),
                scene_description=line_data.get("scene_description"),
            )
            lines.append(line)

        script = Script(
            filename=data.get("filename", "restored"),
            lines=lines,
        )
        return script
    except Exception as e:
        print(f"台本読み込みエラー: {e}")
        return None


def save_prompts_to_output(prompts, output_dir: Path) -> Path | None:
    """プロンプトを出力フォルダに保存"""
    try:
        prompts_path = output_dir / "prompts_backup.json"
        prompts_data = {
            "filename": prompts.filename,
            "prompts": [
                {
                    "number": p.number,
                    "start_time": p.start_time,
                    "end_time": p.end_time,
                    "prompt": p.prompt,
                }
                for p in prompts.prompts
            ],
            "total_images": prompts.total_images,
        }
        with open(prompts_path, "w", encoding="utf-8") as f:
            json.dump(prompts_data, f, ensure_ascii=False, indent=2)
        return prompts_path
    except Exception as e:
        print(f"プロンプト保存エラー: {e}")
        return None


def load_prompts_from_output(output_dir: Path):
    """出力フォルダからプロンプトを読み込み"""
    from src.image.generator import ImagePrompt, ImagePromptList

    prompts_path = output_dir / "prompts_backup.json"
    if not prompts_path.exists():
        return None

    try:
        with open(prompts_path, encoding="utf-8") as f:
            data = json.load(f)

        prompts = [
            ImagePrompt(
                number=p["number"],
                start_time=p["start_time"],
                end_time=p["end_time"],
                prompt=p["prompt"],
            )
            for p in data.get("prompts", [])
        ]

        return ImagePromptList(
            filename=data.get("filename", "restored"),
            prompts=prompts,
        )
    except Exception as e:
        print(f"プロンプト読み込みエラー: {e}")
        return None


def update_history_entry(entry_id: str, updates: dict) -> None:
    """履歴エントリを更新"""
    history = load_generation_history()
    for entry in history:
        if entry["id"] == entry_id:
            for key, value in updates.items():
                if isinstance(value, dict) and key in entry and isinstance(entry[key], dict):
                    entry[key].update(value)
                else:
                    entry[key] = value
            entry["updated_at"] = datetime.now().isoformat()
            break
    save_generation_history(history)


def add_history_entry(entry: dict) -> None:
    """履歴エントリを追加"""
    history = load_generation_history()
    # 同じIDがあれば更新、なければ追加
    existing = next((i for i, e in enumerate(history) if e["id"] == entry["id"]), None)
    if existing is not None:
        history[existing] = entry
    else:
        history.insert(0, entry)  # 新しいものを先頭に
    # 最大50件まで保持
    history = history[:50]
    save_generation_history(history)


def get_history_entry(entry_id: str) -> dict | None:
    """履歴エントリを取得"""
    history = load_generation_history()
    for entry in history:
        if entry["id"] == entry_id:
            return entry
    return None


def delete_history_entry(entry_id: str) -> bool:
    """履歴エントリを削除"""
    history = load_generation_history()
    original_len = len(history)
    history = [e for e in history if e["id"] != entry_id]
    if len(history) < original_len:
        save_generation_history(history)
        return True
    return False


def clear_all_history() -> None:
    """全履歴を削除"""
    save_generation_history([])


def main_page() -> None:
    """P-001: 動画生成メインページ"""
    st.title("🎬 動画生成エージェント")
    st.markdown("台本と画像プロンプトから動画を自動生成します。")

    # セッション状態の初期化
    if "script" not in st.session_state:
        st.session_state.script = None
    if "prompts" not in st.session_state:
        st.session_state.prompts = None
    if "audio_files" not in st.session_state:
        st.session_state.audio_files = {}
    if "generation_complete" not in st.session_state:
        st.session_state.generation_complete = False
    if "output_dir" not in st.session_state:
        st.session_state.output_dir = None
    if "audio_mode" not in st.session_state:
        st.session_state.audio_mode = "batch"  # "batch" or "individual"
    if "output_mode" not in st.session_state:
        st.session_state.output_mode = "自動モード（完成動画出力）"  # デフォルトを自動モードに
    if "output_formats" not in st.session_state:
        st.session_state.output_formats = ["youtube"]  # デフォルト出力形式
    if "script_raw_content" not in st.session_state:
        st.session_state.script_raw_content = ""
    if "reuse_mode" not in st.session_state:
        st.session_state.reuse_mode = {
            "enabled": False,
            "folder": None,
            "audio_files": {},
            "images": {},
            "bgm": None,
            "videos": {},
        }
    if "current_history_id" not in st.session_state:
        st.session_state.current_history_id = None
    if "resume_mode" not in st.session_state:
        st.session_state.resume_mode = {
            "enabled": False,
            "entry": None,
        }
    # 出力フォルダを早期に初期化（履歴読み込みに必要）
    if "custom_output_folder" not in st.session_state:
        st.session_state.custom_output_folder = get_default_output_folder()
    if "pronunciation_suggestions" not in st.session_state:
        st.session_state.pronunciation_suggestions = []
    if "step_mode" not in st.session_state:
        st.session_state.step_mode = True
    if "_last_script_name" not in st.session_state:
        st.session_state._last_script_name = ""
    if "_pending_step" not in st.session_state:
        st.session_state._pending_step = None

    # 履歴セクション（常に表示）
    with st.expander("📜 生成履歴", expanded=True):
        history = load_generation_history()

        if not history:
            st.info("履歴がありません。生成を実行すると履歴が記録されます。")
        else:
            # 中断された生成
            interrupted_entries = [e for e in history if e["status"] == "interrupted"]
            if interrupted_entries:
                st.subheader("⏸️ 中断された生成")
                st.markdown("以下の生成を再開できます。")

                for entry in interrupted_entries[:5]:
                    progress = entry.get("progress", {})
                    completed_steps = sum(1 for v in progress.values() if v)
                    total_steps = len(progress)
                    output_dir_path = Path(entry.get("output_dir", ""))
                    error_log_path = output_dir_path / "error_log.txt"

                    col1, col2, col3, col4, col5 = st.columns([3, 2, 1, 1, 1])
                    with col1:
                        st.markdown(f"**{entry['id']}**")
                        st.caption(f"出力先: {entry.get('output_dir', '不明')}")
                        # エラー情報を表示
                        if entry.get("error"):
                            st.caption(f"❌ エラー: {entry['error'][:50]}...")
                        if error_log_path.exists():
                            st.caption("📋 エラーログあり")
                    with col2:
                        st.progress(completed_steps / total_steps if total_steps > 0 else 0)
                        steps_text = []
                        if progress.get("script_parsed"):
                            steps_text.append("✅台本")
                        if progress.get("audio_generated"):
                            steps_text.append("✅音声")
                        if progress.get("images_generated"):
                            steps_text.append("✅画像")
                        if progress.get("bgm_generated"):
                            steps_text.append("✅BGM")
                        if progress.get("video_generated"):
                            steps_text.append("✅動画")
                        st.caption(" ".join(steps_text) if steps_text else "未開始")
                    with col3:
                        if st.button("▶️ 再開", key=f"resume_{entry['id']}"):
                            output_dir_path = Path(entry.get("output_dir", ""))
                            folder_name = output_dir_path.name

                            # 台本とプロンプトを復元
                            restored_script = load_script_from_output(output_dir_path)
                            restored_prompts = load_prompts_from_output(output_dir_path)

                            if restored_script:
                                st.session_state.script = restored_script
                                st.session_state.resume_mode = {
                                    "enabled": True,
                                    "entry": entry,
                                }

                                if restored_prompts:
                                    st.session_state.prompts = restored_prompts

                                if folder_name:
                                    materials = load_existing_materials(folder_name)
                                    st.session_state.reuse_mode = {
                                        "enabled": True,
                                        "folder": folder_name,
                                        "audio_files": materials["audio_files"],
                                        "images": materials["images"],
                                        "bgm": materials["bgm"],
                                        "videos": materials["videos"],
                                    }

                                # 出力ディレクトリを設定
                                st.session_state.output_dir = output_dir_path

                                st.success(f"✅ {entry['id']} を再開します。台本と素材を復元しました。")
                                st.rerun()
                            else:
                                st.error("❌ 台本ファイルが見つかりません。台本を再度アップロードしてください。")
                                # 素材だけでも読み込む
                                if folder_name:
                                    materials = load_existing_materials(folder_name)
                                    st.session_state.reuse_mode = {
                                        "enabled": True,
                                        "folder": folder_name,
                                        "audio_files": materials["audio_files"],
                                        "images": materials["images"],
                                        "bgm": materials["bgm"],
                                        "videos": materials["videos"],
                                    }
                                    st.info(f"♻️ 素材は読み込みました: 音声{len(materials['audio_files'])}件、画像{len(materials['images'])}枚、動画{len(materials['videos'])}本")
                    with col4:
                        # エラーログ表示ボタン
                        if error_log_path.exists():
                            if st.button("📋 ログ", key=f"log_int_{entry['id']}", help="エラーログを表示"):
                                error_content = read_error_log(output_dir_path)
                                if error_content:
                                    st.code(error_content, language="text")
                    with col5:
                        if st.button("🗑️", key=f"del_int_{entry['id']}", help="この履歴を削除"):
                            delete_history_entry(entry["id"])
                            st.rerun()

                st.divider()

            # 完了した履歴
            completed_entries = [e for e in history if e["status"] == "completed"][:10]
            if completed_entries:
                st.subheader("✅ 完了した生成")

                for entry in completed_entries:
                    col1, col2, col3, col4 = st.columns([4, 1, 1, 1])
                    with col1:
                        st.markdown(f"**{entry['id']}**")
                        st.caption(f"出力先: {entry.get('output_dir', '不明')}")
                        # エラーログの存在を確認
                        folder_path = Path(entry.get("output_dir", ""))
                        error_log_path = folder_path / "error_log.txt"
                        if error_log_path.exists():
                            st.caption("⚠️ エラーログあり")
                    with col2:
                        folder_path = Path(entry.get("output_dir", ""))
                        if folder_path.exists():
                            if st.button("📥 DL", key=f"open_{entry['id']}", help="ダウンロード"):
                                st.session_state[f"show_dl_{entry['id']}"] = True
                        else:
                            st.caption("📁 ファイルなし")
                    with col3:
                        # エラーログ表示ボタン
                        error_log_path = folder_path / "error_log.txt"
                        if error_log_path.exists():
                            if st.button("📋 ログ", key=f"log_{entry['id']}", help="エラーログを表示"):
                                error_content = read_error_log(folder_path)
                                if error_content:
                                    st.code(error_content, language="text")
                    with col4:
                        if st.button("🗑️", key=f"del_comp_{entry['id']}", help="この履歴を削除"):
                            delete_history_entry(entry["id"])
                            st.rerun()

                    # ダウンロードパネル展開（準備ボタンで1ファイルだけ読み込み）
                    if st.session_state.get(f"show_dl_{entry['id']}") and folder_path.exists():
                        dl_files = sorted([f for f in folder_path.rglob("*") if f.is_file() and f.name != "error_log.txt"])
                        if dl_files:
                            h_file_options = {}
                            for f in dl_files:
                                f_mb = f.stat().st_size / (1024 * 1024)
                                size_label = f"{f_mb:.0f}MB" if f_mb >= 1 else f"{f_mb * 1024:.0f}KB"
                                rel_path = f.relative_to(folder_path)
                                h_file_options[f"{rel_path} ({size_label})"] = str(f)
                            h_selected = st.selectbox(
                                "ファイルを選択",
                                options=list(h_file_options.keys()),
                                key=f"hsel_{entry['id']}",
                            )
                            h_ready_key = f"hready_{entry['id']}"
                            if st.button("📦 ダウンロード準備", key=f"hprep_{entry['id']}"):
                                st.session_state[h_ready_key] = h_file_options.get(h_selected, "")
                            if st.session_state.get(h_ready_key) and h_selected and h_file_options.get(h_selected) == st.session_state.get(h_ready_key):
                                h_ready_file = Path(st.session_state[h_ready_key])
                                if h_ready_file.exists():
                                    st.download_button(
                                        label="📥 ダウンロード",
                                        data=h_ready_file.read_bytes(),
                                        file_name=h_ready_file.name,
                                        mime="application/octet-stream",
                                        key=f"hdl_{entry['id']}",
                                    )
                        else:
                            st.caption("ダウンロード可能なファイルがありません")

            # 全削除ボタン
            st.divider()
            col1, col2 = st.columns([3, 1])
            with col2:
                if st.button("🗑️ 全履歴を削除", type="secondary"):
                    clear_all_history()
                    st.success("✅ 履歴を全て削除しました")
                    st.rerun()

    # 素材再利用オプション（STEP 0）
    existing_folders = get_existing_output_folders()  # list of (name, path) tuples
    with st.expander("♻️ 素材再利用（オプション）", expanded=False):
        st.markdown("複数の案件から素材を選んで組み合わせできます。APIクレジットを節約できます。")

        if existing_folders:
            # フォルダ選択肢を作成（表示名: パス）
            folder_options = {f"{name} ({path})": path for name, path in existing_folders}
            folder_display_names = ["新規生成"] + list(folder_options.keys())

            # 素材ごとに個別選択
            st.markdown("**素材ごとに選択:**")

            col_audio, col_image, col_bgm, col_video = st.columns(4)

            with col_audio:
                st.markdown("🎤 **音声**")
                audio_source = st.selectbox(
                    "音声の取得元",
                    options=folder_display_names,
                    key="reuse_audio_source",
                    label_visibility="collapsed",
                )

            with col_image:
                st.markdown("🖼️ **画像**")
                image_source = st.selectbox(
                    "画像の取得元",
                    options=folder_display_names,
                    key="reuse_image_source",
                    label_visibility="collapsed",
                )

            with col_bgm:
                st.markdown("🎵 **BGM**")
                bgm_source = st.selectbox(
                    "BGMの取得元",
                    options=folder_display_names,
                    key="reuse_bgm_source",
                    label_visibility="collapsed",
                )

            with col_video:
                st.markdown("🎬 **動画**")
                video_source = st.selectbox(
                    "動画の取得元",
                    options=folder_display_names,
                    key="reuse_video_source",
                    label_visibility="collapsed",
                )

            # 読み込みボタン
            if st.button("📂 選択した素材を読み込む", type="secondary"):
                audio_files = {}
                images = {}
                bgm = None
                videos = {}

                # 音声を読み込み
                if audio_source != "新規生成":
                    audio_path = folder_options[audio_source]
                    materials = load_existing_materials(audio_path)
                    audio_files = materials["audio_files"]

                # 画像を読み込み
                if image_source != "新規生成":
                    image_path = folder_options[image_source]
                    materials = load_existing_materials(image_path)
                    images = materials["images"]

                # BGMを読み込み
                if bgm_source != "新規生成":
                    bgm_path = folder_options[bgm_source]
                    materials = load_existing_materials(bgm_path)
                    bgm = materials["bgm"]

                # 動画を読み込み
                if video_source != "新規生成":
                    video_path = folder_options[video_source]
                    materials = load_existing_materials(video_path)
                    videos = materials["videos"]

                # 何か選択されていれば再利用モードを有効化
                has_materials = audio_files or images or bgm or videos
                st.session_state.reuse_mode = {
                    "enabled": has_materials,
                    "folder": None,  # 複数ソースのため単一フォルダは指定しない
                    "audio_files": audio_files,
                    "images": images,
                    "bgm": bgm,
                    "videos": videos,
                }

                if has_materials:
                    st.success("✅ 素材を読み込みました")
                else:
                    st.info("ℹ️ 全て新規生成が選択されています")
                st.rerun()

        else:
            st.info("📁 再利用可能なフォルダがありません。生成を実行すると、ここに表示されます。")

        # 読み込み結果を表示
        if st.session_state.reuse_mode["enabled"]:
            st.divider()
            st.markdown("**読み込み済み素材:**")
            col1, col2, col3, col4 = st.columns(4)
            with col1:
                audio_count = len(st.session_state.reuse_mode["audio_files"])
                st.metric("🎤 音声", f"{audio_count}件")
            with col2:
                image_count = len(st.session_state.reuse_mode["images"])
                st.metric("🖼️ 画像", f"{image_count}枚")
            with col3:
                bgm_status = "あり" if st.session_state.reuse_mode["bgm"] else "なし"
                st.metric("🎵 BGM", bgm_status)
            with col4:
                video_count = len(st.session_state.reuse_mode.get("videos", {}))
                st.metric("🎬 動画", f"{video_count}本")

            if st.button("❌ 再利用モードを解除"):
                st.session_state.reuse_mode = {
                    "enabled": False,
                    "folder": None,
                    "audio_files": {},
                    "images": {},
                    "bgm": None,
                    "videos": {},
                }
                st.rerun()

    # STEP 1: ファイルアップロード
    st.header("STEP 1: ファイルアップロード")

    col1, col2 = st.columns(2)

    with col1:
        st.subheader("📝 台本ファイル")
        script_file = st.file_uploader(
            "Word(.docx)またはテキスト(.txt)ファイルをアップロード",
            type=["docx", "txt"],
            key="script_file",
        )
        if script_file:
            st.success(f"✅ {script_file.name} をアップロードしました")

            # 新しいファイルの場合のみパース（毎回のリランでは実行しない）
            last_script_name = st.session_state.get("_last_script_name", "")
            if script_file.name != last_script_name:
                st.session_state._last_script_name = script_file.name
                # 生のコンテンツを保存（項数検出用）
                if script_file.name.lower().endswith(".docx"):
                    from docx import Document
                    doc = Document(BytesIO(script_file.getvalue()))
                    st.session_state.script_raw_content = "\n".join(para.text for para in doc.paragraphs)
                    script_file.seek(0)  # ファイルポインタをリセット
                else:
                    st.session_state.script_raw_content = script_file.getvalue().decode("utf-8")
                    script_file.seek(0)  # ファイルポインタをリセット
                # 台本をパース
                parser = ScriptParser()
                st.session_state.script = parser.parse_uploaded_file(script_file)

                # 読み確認を自動実行（high/medium を自動適用）
                try:
                    from src.audio.pronunciation import apply_suggestions, check_pronunciation

                    full_text = "\n".join(
                        line.text for line in st.session_state.script.lines
                    )
                    with st.spinner("読み確認を実行中..."):
                        suggestions = check_pronunciation(full_text)
                    if suggestions:
                        auto_indices = [
                            i for i, s in enumerate(suggestions)
                            if s.confidence in ("high", "medium")
                        ]
                        if auto_indices:
                            apply_suggestions(
                                st.session_state.script.lines,
                                suggestions,
                                auto_indices,
                            )
                            st.info(f"読み仮名を{len(auto_indices)}件自動適用しました")
                        # low confidence は手動確認用に残す
                        remaining = [
                            s for s in suggestions if s.confidence == "low"
                        ]
                        st.session_state.pronunciation_suggestions = remaining
                except Exception as pron_err:
                    st.warning(f"⚠️ 読み自動チェック: {pron_err}")

    with col2:
        st.subheader("🖼️ 画像プロンプトファイル")
        prompt_file = st.file_uploader(
            "Word(.docx)またはテキスト(.txt)ファイルをアップロード（任意）",
            type=["docx", "txt"],
            key="prompt_file",
        )
        if prompt_file:
            st.success(f"✅ {prompt_file.name} をアップロードしました")

            # ファイル内容を読み取り
            if prompt_file.name.lower().endswith(".docx"):
                from docx import Document
                doc = Document(BytesIO(prompt_file.getvalue()))
                file_content = "\n".join(para.text for para in doc.paragraphs)
            else:
                file_content = prompt_file.getvalue().decode("utf-8")

            # プロンプトファイルの生コンテンツを保存（チャプター検出用）
            st.session_state.prompt_raw_content = file_content

            # プロンプトをパース（既に読み取ったfile_contentを直接渡す）
            generator = ImageGenerator()
            parsed_prompts = generator.parse_prompt_text(file_content, prompt_file.name)
            st.session_state.prompts = parsed_prompts

            # パース結果をすぐに表示（デバッグ用）
            if parsed_prompts.total_images > 0:
                st.info(f"🎨 {parsed_prompts.total_images}件のプロンプトを検出しました")
            else:
                st.warning(
                    "⚠️ 画像プロンプトを検出できませんでした。\n\n"
                    "**対応形式:**\n"
                    "- `[番号] 開始時間-終了時間 | プロンプト` 形式\n"
                    "  例: `[1] 0:00-0:15 | 青空の下で微笑む人物`\n"
                    "- `【画像生成プロンプト】` セクション形式\n"
                    "  例: `1. タイトル\\n【画像生成プロンプト】\\nプロンプト内容`\n"
                    "- `番号. プロンプト` シンプルリスト形式\n"
                    "  例: `1. 青空の下で微笑む人物`"
                )
                # ファイル内容のプレビューを表示（デバッグ用）
                with st.expander("📄 ファイル内容プレビュー（先頭500文字）"):
                    st.code(file_content[:500] if len(file_content) > 500 else file_content)
        elif st.session_state.script and not st.session_state.prompts:
            st.info("💡 画像プロンプトファイルがない場合、台本から自動生成できます")

    # STEP 2: 台本プレビュー（ファイルアップロード後に表示）
    script = st.session_state.script
    if script:
        st.header("STEP 2: 台本プレビュー＆前処理")

        st.info(f"📄 ファイル: {script.filename} | セリフ数: {script.total_lines}")

        # セリフ一覧を表示
        for line in script.lines:
            col1, col2 = st.columns([1, 4])
            with col1:
                speaker_label = "🔵 Speaker1" if line.speaker == "speaker1" else "🟠 Speaker2"
                st.markdown(f"**{line.number}. {speaker_label}**")
            with col2:
                # 情景補足があれば表示
                if line.scene_description:
                    st.markdown(f"~~({line.scene_description})~~ *（除去済み）*")
                st.markdown(line.text)

        st.markdown("""
        **自動前処理:**
        - `(...)` 形式の情景補足は自動除去されます
        - `{漢字|読み}` 形式で読み仮名を指定できます
        """)

        # 読み確認（自動適用済み + 残りの低確信度候補）
        st.divider()
        st.markdown("🔍 **読み確認**: high/medium 確信度の修正は台本アップロード時に自動適用済みです。")

        if st.session_state.pronunciation_suggestions:
            suggestions = st.session_state.pronunciation_suggestions
            st.markdown(f"**確認が必要な候補: {len(suggestions)}件**（低確信度）")

            selected = []
            for i, s in enumerate(suggestions):
                conf_icon = {"high": "🔴", "medium": "🟡", "low": "🟢"}.get(s.confidence, "⚪")
                checked = st.checkbox(
                    f"{conf_icon} [{s.confidence}] {s.original} → {s.notation}",
                    value=False,
                    key=f"pron_{i}",
                )
                if checked:
                    selected.append(i)

            if selected and st.button("✅ 選択した修正を適用", key="apply_pronunciation_btn"):
                from src.audio.pronunciation import apply_suggestions

                apply_suggestions(script.lines, suggestions, selected)
                st.session_state.pronunciation_suggestions = []
                st.success(f"✅ {len(selected)}件の修正を台本に適用しました")
                st.rerun()

    # 画像プロンプト自動生成オプション
    if script and not st.session_state.prompts:
        st.subheader("🖼️ 画像プロンプト自動生成")
        st.markdown("台本の内容からAIが自動的に画像プロンプトを生成します。")

        # 台本から項数を自動検出（生コンテンツから）
        raw_content = st.session_state.get("script_raw_content", "")
        if raw_content:
            detected_items = count_script_items_from_content(raw_content)
        else:
            detected_items = count_script_items(script)

        if detected_items > 0:
            st.info(f"📊 台本から検出された項数: {detected_items}")
        else:
            detected_items = script.total_lines
            st.info(f"📊 項番号が検出されませんでした。行数を使用: {detected_items}")

        # 画像枚数の設定
        # セッションに保存された値があれば使用、なければ検出値を使用
        default_num_images = st.session_state.get("user_specified_num_images", min(detected_items, 100))
        num_images = st.number_input(
            "生成する画像の枚数",
            min_value=1,
            max_value=100,
            value=default_num_images,
            help=f"台本から{detected_items}項を検出しました。必要に応じて調整してください。",
            key="num_images_input"
        )
        # ユーザーが指定した枚数をセッションに保存（生成時に使用）
        st.session_state["user_specified_num_images"] = num_images

        if st.button("🎨 台本から画像プロンプトを自動生成", type="primary"):
            with st.spinner("AIが台本を分析して画像プロンプトを生成中..."):
                try:
                    auto_prompts = generate_image_prompts_from_script(script, num_images)
                    st.session_state.prompts = auto_prompts
                    st.success(f"✅ {auto_prompts.total_images}件の画像プロンプトを生成しました")
                    st.rerun()
                except Exception as e:
                    st.error(f"❌ 画像プロンプト生成エラー: {e}")

    # 画像プロンプトのプレビュー
    prompts = st.session_state.prompts
    st.subheader("🖼️ 画像プロンプト一覧")
    if prompts and prompts.total_images > 0:
        st.success(f"📄 ファイル: {prompts.filename} | ✅ 画像数: {prompts.total_images}")
        for p in prompts.prompts:
            st.markdown(f"**[{p.number}]** `{p.start_time}` - `{p.end_time}` | {p.prompt}")
    else:
        st.warning("⚠️ 画像プロンプトがまだ生成されていません。上の「🎨 台本から画像プロンプトを自動生成」ボタンをクリックしてください。")
        st.info("💡 または、生成開始時に自動的に生成されます。")

    # STEP 3: 音声プレビュー
    if script:
        st.header("STEP 3: 音声プレビュー＆確認")

        # APIキー確認
        has_google_creds = bool(get_gcp_credentials())

        if not has_google_creds:
            st.warning("⚠️ Google Cloud TTSのAPIキーが必要です。設定ページで設定してください。")
        else:
            st.success("✅ Google Cloud TTS APIキー設定済み")

            # --- 話者チェックテーブル ---
            st.subheader("話者チェック")

            import pandas as pd

            speakers_in_script = sorted({line.speaker for line in script.lines})

            # 話者変更用セレクトボックス（expanderで折りたたみ）
            speaker_changed = False
            with st.expander("話者の割り当てを変更", expanded=False):
                for idx, line in enumerate(script.lines):
                    col1, col2 = st.columns([1, 4])
                    with col1:
                        new_speaker = st.selectbox(
                            f"No.{line.number}",
                            options=speakers_in_script,
                            index=speakers_in_script.index(line.speaker),
                            key=f"speaker_select_{idx}",
                        )
                    with col2:
                        st.text(line.text[:60])
                    if new_speaker != line.speaker:
                        script.lines[idx].speaker = new_speaker
                        speaker_changed = True

            if speaker_changed:
                st.session_state.script = script
                st.rerun()

            # 一覧テーブル表示
            table_data = []
            for line in script.lines:
                table_data.append({
                    "No.": line.number,
                    "話者": line.speaker,
                    "セリフ冒頭": line.text[:50],
                })
            df = pd.DataFrame(table_data)

            def highlight_speaker(row: pd.Series) -> list[str]:
                if row["話者"] == "speaker1":
                    return ["background-color: #e6f3ff"] * len(row)
                elif row["話者"] == "speaker2":
                    return ["background-color: #fff3e6"] * len(row)
                return [""] * len(row)

            st.dataframe(
                df.style.apply(highlight_speaker, axis=1),
                use_container_width=True,
                hide_index=True,
            )

            # 同一話者3回以上連続の警告
            consecutive_count = 1
            warnings: list[str] = []
            for i in range(1, len(script.lines)):
                if script.lines[i].speaker == script.lines[i - 1].speaker:
                    consecutive_count += 1
                    if consecutive_count == 3:
                        warnings.append(
                            f"No.{script.lines[i - 2].number}〜{script.lines[i].number}: "
                            f"{script.lines[i].speaker} が3回以上連続しています"
                        )
                else:
                    consecutive_count = 1
            if warnings:
                for w in warnings:
                    st.warning(f"⚠️ {w}")

            st.divider()

            # 個別プレビュー
            selected_line = st.selectbox(
                "プレビューするセリフを選択",
                options=range(len(script.lines)),
                format_func=lambda i: f"{script.lines[i].number}. {script.lines[i].speaker}: {script.lines[i].text[:30]}...",
            )

            if st.button("🎤 選択したセリフをプレビュー", type="secondary"):
                line = script.lines[selected_line]
                try:
                    with st.spinner("音声を生成中..."):
                        tts = TTSClient()
                        temp_path = Path("temp") / f"preview_{line.number}.wav"
                        temp_path.parent.mkdir(exist_ok=True)
                        wav_path = tts.synthesize(line.text, line.speaker, temp_path)

                        st.audio(str(wav_path), format="audio/wav")
                        st.session_state.audio_files[line.number] = str(wav_path)
                except Exception as e:
                    st.error(f"❌ 音声生成エラー: {e}")

            # 音声生成モード選択
            audio_mode_options = ["一括生成（1本のファイル・推奨）", "個別生成（セリフごとのファイル）"]
            default_index = 0 if st.session_state.audio_mode == "batch" else 1
            audio_mode = st.radio(
                "音声生成モード",
                audio_mode_options,
                index=default_index,
                horizontal=True,
                help="一括生成: マルチスピーカーで自然な会話を1つのファイルに。個別生成: 各セリフを別々のファイルに。"
            )
            # セッションステートに保存
            st.session_state.audio_mode = "batch" if audio_mode == audio_mode_options[0] else "individual"

            if st.button("🔊 全セリフの音声を生成", type="primary"):
                progress = st.progress(0)
                status = st.empty()

                try:
                    tts = TTSClient()
                    output_dir = get_output_dir()
                    audio_dir = output_dir / "audio"
                    audio_dir.mkdir(exist_ok=True)

                    if st.session_state.audio_mode == "batch":
                        # マルチスピーカー一括生成
                        def update_progress(current, total, message):
                            """進捗を更新するコールバック"""
                            progress.progress((current + 1) / total)
                            status.text(f"🎤 生成中: {current + 1}/{total} - {message}")

                        status.text("🎤 マルチスピーカー音声を一括生成中...")
                        output_path = audio_dir / "full_audio.wav"
                        wav_path = tts.synthesize_script(script, output_path, progress_callback=update_progress)
                        st.session_state.audio_files["full"] = str(wav_path)
                        progress.progress(1.0)
                        st.session_state.output_dir = output_dir
                        st.success(f"✅ 音声を1本のファイルに生成しました: {wav_path.name}")
                        st.audio(str(wav_path), format="audio/wav")
                    else:
                        # 個別生成（従来方式）
                        for i, line in enumerate(script.lines):
                            status.text(f"生成中: {i + 1}/{script.total_lines} - {line.speaker}")
                            output_path = audio_dir / f"{line.number:03d}_{line.speaker}.wav"
                            tts.synthesize(line.text, line.speaker, output_path)
                            st.session_state.audio_files[line.number] = str(output_path)
                            progress.progress((i + 1) / script.total_lines)

                        st.session_state.output_dir = output_dir
                        st.success(f"✅ {script.total_lines}件の音声を生成しました")
                except Exception as e:
                    st.error(f"❌ 音声生成エラー: {e}")

    # STEP 4: モード選択＆生成実行
    if script and prompts:
        st.header("STEP 4: モード選択＆生成実行")

        mode_options = ["Filmoraモード（素材出力）", "自動モード（完成動画出力）"]
        default_mode_index = 1 if st.session_state.output_mode == "自動モード（完成動画出力）" else 0
        mode = st.radio(
            "出力モードを選択",
            mode_options,
            index=default_mode_index,
            horizontal=True,
        )
        st.session_state.output_mode = mode

        # 自動モードでは一括生成（全ステップ連続実行）を強制
        if mode == "自動モード（完成動画出力）":
            st.session_state.step_mode = False
            st.session_state.audio_mode = "batch"

        output_formats = []
        if mode == "自動モード（完成動画出力）":
            st.subheader("出力形式を選択")
            output_formats = st.multiselect(
                "出力する形式を選択してください（複数選択可）",
                ["youtube", "instagram_reel", "instagram_feed", "tiktok"],
                default=st.session_state.output_formats,
                format_func=lambda x: {
                    "youtube": "YouTube (1920×1080)",
                    "instagram_reel": "Instagram リール (1080×1920)",
                    "instagram_feed": "Instagram フィード (1080×1080)",
                    "tiktok": "TikTok (1080×1920)",
                }.get(x, x),
            )
            st.session_state.output_formats = output_formats

            # 出力形式が選択されていない場合の警告
            if not output_formats:
                st.warning("⚠️ 出力形式を1つ以上選択してください")

        st.divider()

        # 出力フォルダ設定
        with st.expander("📁 出力フォルダ設定", expanded=False):
            import os
            settings = load_settings()
            default_output = settings.get("defaults", {}).get("output_folder", "output")

            # ホームディレクトリとよく使うパスを取得
            home_dir = os.path.expanduser("~")

            # OS別のデフォルトフォルダを取得
            default_local_folder = get_default_output_folder()

            preset_paths = {
                "推奨（ローカル保存）": default_local_folder,
                "ドキュメント": os.path.join(home_dir, "Documents"),
                "デスクトップ": os.path.join(home_dir, "Desktop"),
                "ダウンロード": os.path.join(home_dir, "Downloads"),
                "ホーム": home_dir,
                "相対パス (output)": "output",
                "カスタム入力": "_custom_",
            }

            # プリセット選択（推奨をデフォルトに）
            selected_preset = st.selectbox(
                "出力先を選択",
                options=list(preset_paths.keys()),
                index=0,  # 「推奨（ローカル保存）」がデフォルト
                key="output_preset_select",
            )

            if selected_preset == "カスタム入力":
                # カスタムパス入力
                custom_output = st.text_input(
                    "カスタムパスを入力",
                    value=st.session_state.get("custom_output_folder", default_output),
                    help="絶対パスまたは相対パスで指定できます。"
                )
            else:
                custom_output = preset_paths[selected_preset]

            st.session_state.custom_output_folder = custom_output

            st.info(f"📂 現在の出力先: `{custom_output}/[タイムスタンプ]/`")

            # フォルダが存在するか確認
            if os.path.isabs(custom_output) and not os.path.exists(custom_output):
                st.warning(f"⚠️ フォルダが存在しません。生成時に自動作成されます。")

        st.divider()

        # API設定状況確認
        api_status = {
            "Google Cloud TTS": bool(get_gcp_credentials()),
            "Gemini API": bool(get_env_var("GOOGLE_API_KEY")),
            "Beatoven.ai": bool(get_env_var("BEATOVEN_API_KEY")),
            "Pexels": bool(get_env_var("PEXELS_API_KEY")),
        }

        with st.expander("📋 API設定状況"):
            for name, is_set in api_status.items():
                status = "✅ 設定済み" if is_set else "❌ 未設定"
                st.text(f"{name}: {status}")

        # 生成モード切替
        st.subheader("🔧 生成モード")
        step_mode = st.toggle(
            "ステップモード（推奨: 中断防止）",
            value=st.session_state.step_mode,
            help="ONにすると各ステップを個別ボタンで実行。メモリ解放が行われ中断を防止します。",
            key="step_mode_toggle",
        )
        st.session_state.step_mode = step_mode

        if step_mode:
            # ステップモード: 個別ボタンUI
            st.markdown("各ステップを個別に実行できます。完了後にメモリが解放されます。")

            # 出力ディレクトリ準備
            if not st.session_state.output_dir:
                st.session_state.output_dir = get_output_dir()
            step_output_dir = Path(st.session_state.output_dir)

            # ステップ状態検出
            step_status = detect_step_status(step_output_dir)

            # 履歴エントリ準備
            step_history = None
            try:
                step_history = _get_or_create_history_entry(step_output_dir)
                step_history["settings"]["output_mode"] = mode
                step_history["settings"]["output_formats"] = output_formats
                add_history_entry(step_history)
            except Exception:
                pass

            # 台本・プロンプトを保存
            if step_history:
                step_history["progress"]["script_parsed"] = True
                script_file = save_script_to_output(script, step_output_dir)
                if script_file:
                    step_history["files"]["script_file"] = str(script_file)
                if prompts:
                    prompts_file = save_prompts_to_output(prompts, step_output_dir)
                    if prompts_file:
                        step_history["files"]["prompts_file"] = str(prompts_file)
                add_history_entry(step_history)

            # --- 実行予約の処理（ボタンではなくフラグで駆動） ---
            pending_step = st.session_state.get("_pending_step")
            if pending_step:
                st.session_state._pending_step = None
                step_names = {"audio": "音声", "bgm": "BGM", "bg_video": "背景動画", "images": "画像", "timeline": "タイムライン"}
                st.info(f"⏳ {step_names.get(pending_step, pending_step)} の生成を開始...")
                _step_ok = False
                try:
                    if pending_step == "audio":
                        result = run_step_audio(script, step_output_dir, step_history)
                        st.write(f"結果: 成功={result.get('success')}, ファイル数={len(result.get('files', {}))}, エラー={result.get('error')}")
                    elif pending_step == "bgm":
                        run_step_bgm(script, prompts, step_output_dir, step_history)
                    elif pending_step == "bg_video":
                        run_step_bg_video(script, prompts, step_output_dir, step_history)
                    elif pending_step == "images":
                        run_step_images(script, prompts, step_output_dir, step_history)
                    elif pending_step == "timeline":
                        materials = load_existing_materials(str(step_output_dir))
                        st.session_state.audio_files = materials["audio_files"]
                        run_step_timeline(
                            script, prompts, mode, output_formats, step_output_dir,
                            materials["images"], materials["videos"],
                            Path(materials["bgm"]) if materials["bgm"] else None,
                            step_history,
                        )
                        st.session_state.generation_complete = True
                    _step_ok = True
                except Exception as e:
                    st.error(f"生成エラー: {e}")
                    import traceback
                    st.code(traceback.format_exc())
                if _step_ok:
                    st.rerun()

            # --- ステップダッシュボード（ボタンはフラグを立てるだけ） ---

            # STEP 1: 音声
            s1_icon = "✅" if step_status["audio"] else "⬜"
            audio_count = len(list((step_output_dir / "audio").glob("*.wav"))) if (step_output_dir / "audio").exists() else 0
            st.markdown(f"### {s1_icon} STEP 1: 音声生成 ({audio_count}件)")
            s1_col1, s1_col2 = st.columns([3, 1])
            with s1_col2:
                s1_label = "🔄 再生成" if step_status["audio"] else "▶️ 生成開始"
                if st.button(s1_label, key="step_audio_btn", use_container_width=True):
                    st.session_state._pending_step = "audio"
                    st.rerun()

            # STEP 2: BGM
            s2_icon = "✅" if step_status["bgm"] else "⬜"
            st.markdown(f"### {s2_icon} STEP 2: BGM生成")
            s2_col1, s2_col2 = st.columns([3, 1])
            with s2_col2:
                s2_label = "🔄 再生成" if step_status["bgm"] else "▶️ 生成開始"
                if st.button(s2_label, key="step_bgm_btn", use_container_width=True):
                    st.session_state._pending_step = "bgm"
                    st.rerun()

            # STEP 3: 背景動画
            s3_icon = "✅" if step_status["bg_video"] else "⬜"
            st.markdown(f"### {s3_icon} STEP 3: 背景動画取得")
            s3_col1, s3_col2 = st.columns([3, 1])
            with s3_col2:
                s3_label = "🔄 再生成" if step_status["bg_video"] else "▶️ 生成開始"
                if st.button(s3_label, key="step_bg_video_btn", use_container_width=True):
                    st.session_state._pending_step = "bg_video"
                    st.rerun()

            # STEP 4: 画像
            s4_icon = "✅" if step_status["images"] else "⬜"
            img_count = 0
            img_dir = step_output_dir / "images"
            if img_dir.exists():
                img_count = len(list(img_dir.glob("*.png"))) + len(list(img_dir.glob("*.jpg")))
            st.markdown(f"### {s4_icon} STEP 4: 画像生成 ({img_count}枚)")
            s4_col1, s4_col2 = st.columns([3, 1])
            with s4_col2:
                s4_label = "🔄 再生成" if step_status["images"] else "▶️ 生成開始"
                if st.button(s4_label, key="step_images_btn", use_container_width=True):
                    st.session_state._pending_step = "images"
                    st.rerun()

            # STEP 5: タイムライン/動画合成
            s5_icon = "✅" if step_status["timeline"] else "⬜"
            st.markdown(f"### {s5_icon} STEP 5: タイムライン/動画合成")
            s5_col1, s5_col2 = st.columns([3, 1])
            with s5_col2:
                s5_label = "🔄 再生成" if step_status["timeline"] else "▶️ 生成開始"
                if st.button(s5_label, key="step_timeline_btn", use_container_width=True):
                    if step_status["timeline"]:
                        _clear_step_files(step_output_dir, ["timeline.csv"])
                        _clear_step_files(step_output_dir / "videos", ["*.mp4"], exclude_subdir="backgrounds")
                    st.session_state._pending_step = "timeline"
                    st.rerun()

            st.info(f"📂 出力先: `{step_output_dir}`")

        else:
            # 一括モード（従来方式）
            st.subheader("🎯 生成する素材を選択")
            col_audio, col_image, col_bgm, col_bg_video = st.columns(4)
            with col_audio:
                generate_audio = st.checkbox("🎤 音声", value=True, key="generate_audio_checkbox")
            with col_image:
                generate_images = st.checkbox("🖼️ 画像", value=True, key="generate_images_checkbox")
            with col_bgm:
                generate_bgm = st.checkbox("🎵 BGM", value=True, key="generate_bgm_checkbox")
            with col_bg_video:
                generate_bg_video = st.checkbox("🎬 背景動画", value=True, key="generate_bg_video_checkbox")

            if not generate_audio and not generate_images and not generate_bgm and not generate_bg_video:
                st.warning("⚠️ 少なくとも1つの素材を選択してください")

            st.divider()

            if st.button("🚀 生成を開始", type="primary", use_container_width=True):
                if not generate_audio and not generate_images and not generate_bgm and not generate_bg_video:
                    st.error("❌ 少なくとも1つの素材を選択してください")
                elif not all(api_status.values()):
                    st.warning("⚠️ 一部のAPIキーが未設定です。設定ページで設定してください。")
                elif mode == "自動モード（完成動画出力）" and not output_formats:
                    st.error("❌ 出力形式を1つ以上選択してください")
                else:
                    run_generation(script, prompts, mode, output_formats, generate_audio, generate_images, generate_bgm, generate_bg_video)

    # STEP 5: 結果ダウンロード
    st.header("STEP 5: 結果ダウンロード")

    # 出力ディレクトリの確認（完了・失敗に関わらず）
    output_dir = None
    if st.session_state.output_dir:
        output_dir = Path(st.session_state.output_dir)

    if output_dir and output_dir.exists():
        # ファイル数をカウント
        all_files = list(output_dir.rglob("*"))
        file_count = len([f for f in all_files if f.is_file()])

        if file_count > 0:
            # セッション状態と生成履歴の両方からステータスを判定
            is_complete = st.session_state.generation_complete
            if not is_complete:
                # セッションがリセットされた場合、履歴から復元
                history = load_generation_history()
                for entry in reversed(history):
                    if entry.get("output_dir") == str(output_dir):
                        if entry.get("status") == "completed":
                            is_complete = True
                            st.session_state.generation_complete = True
                        break

            if is_complete:
                st.success(f"✅ 生成完了！出力先: {output_dir}")
            else:
                st.warning(f"⚠️ 生成が中断されましたが、一部の素材は保存されています。出力先: {output_dir}")

            # ファイル種別ごとのカウント
            audio_files = list((output_dir / "audio").rglob("*")) if (output_dir / "audio").exists() else []
            image_files = list((output_dir / "images").rglob("*")) if (output_dir / "images").exists() else []
            bgm_files = list((output_dir / "bgm").rglob("*")) if (output_dir / "bgm").exists() else []
            # 完成動画（videos直下のmp4）と素材動画（backgrounds内）を分離
            videos_dir = output_dir / "videos"
            final_videos = [f for f in videos_dir.glob("*.mp4") if f.is_file()] if videos_dir.exists() else []
            bg_videos = list((videos_dir / "backgrounds").rglob("*.mp4")) if (videos_dir / "backgrounds").exists() else []

            # 完成動画がある場合は目立たせて表示
            if final_videos:
                st.markdown("### 🎬 完成動画")
                for fv in final_videos:
                    fv_size_mb = fv.stat().st_size / (1024 * 1024)
                    fcol1, fcol2 = st.columns([3, 1])
                    with fcol1:
                        st.markdown(f"**{fv.name}** ({fv_size_mb:.0f}MB)")
                    with fcol2:
                        with open(fv, "rb") as vf:
                            st.download_button(
                                label=f"📥 ダウンロード",
                                data=vf,
                                file_name=fv.name,
                                mime="video/mp4",
                                key=f"dl_final_{fv.name}",
                            )
                st.divider()

            col1, col2, col3, col4 = st.columns(4)
            with col1:
                st.metric("🎤 音声", f"{len([f for f in audio_files if f.is_file()])}件")
            with col2:
                st.metric("🖼️ 画像", f"{len([f for f in image_files if f.is_file()])}枚")
            with col3:
                st.metric("🎵 BGM", f"{len([f for f in bgm_files if f.is_file()])}件")
            with col4:
                st.metric("🎬 素材動画", f"{len(bg_videos)}本")

            # カテゴリ別ZIPダウンロード
            def _create_zip(files: list, base_dir: Path) -> BytesIO:
                """指定ファイルリストからZIPバッファを作成"""
                buf = BytesIO()
                with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
                    for f in files:
                        if f.is_file():
                            zf.write(f, f.relative_to(base_dir))
                buf.seek(0)
                return buf

            def _dir_size_mb(files: list) -> float:
                """ファイルリストの合計サイズ(MB)"""
                return sum(f.stat().st_size for f in files if f.is_file()) / (1024 * 1024)

            categories = [
                ("🎤 音声", "audio", [f for f in audio_files if f.is_file()]),
                ("🖼️ 画像", "images", [f for f in image_files if f.is_file()]),
                ("🎵 BGM", "bgm", [f for f in bgm_files if f.is_file()]),
                ("🎬 素材動画", "videos", bg_videos),
            ]

            MAX_ZIP_MB = 300  # ZIPダウンロード上限(MB)

            st.markdown("**カテゴリ別ダウンロード**")
            dl_cols = st.columns(len(categories))
            for col, (label, folder_name, files) in zip(dl_cols, categories):
                with col:
                    if files:
                        size_mb = _dir_size_mb(files)
                        if size_mb <= MAX_ZIP_MB:
                            zip_buf = _create_zip(files, output_dir)
                            st.download_button(
                                label=f"{label} ({size_mb:.0f}MB)",
                                data=zip_buf,
                                file_name=f"{folder_name}_{output_dir.name}.zip",
                                mime="application/zip",
                                key=f"dl_{folder_name}",
                            )
                        else:
                            # 大容量: 選択→準備→ダウンロードの3段階（メモリ節約）
                            file_options = {f"{f.name} ({f.stat().st_size / (1024*1024):.0f}MB)": str(f) for f in files}
                            selected = st.selectbox(
                                f"{label} ({size_mb:.0f}MB) - ファイルを選択",
                                options=list(file_options.keys()),
                                key=f"sel_{folder_name}",
                            )
                            ready_key = f"ready_{folder_name}"
                            if st.button("📦 ダウンロード準備", key=f"prep_{folder_name}"):
                                st.session_state[ready_key] = file_options.get(selected, "")
                            if st.session_state.get(ready_key) and selected and file_options.get(selected) == st.session_state.get(ready_key):
                                ready_file = Path(st.session_state[ready_key])
                                if ready_file.exists():
                                    st.download_button(
                                        label=f"📥 ダウンロード",
                                        data=ready_file.read_bytes(),
                                        file_name=ready_file.name,
                                        mime="application/octet-stream",
                                        key=f"dl_{folder_name}_sel",
                                    )
                    else:
                        st.button(f"{label} なし", disabled=True, key=f"dl_{folder_name}")

            # 一括ダウンロード（500MB以下の場合のみ）
            all_files = [f for f in output_dir.rglob("*") if f.is_file()]
            total_mb = _dir_size_mb(all_files)
            if total_mb <= 500:
                zip_buffer = _create_zip(all_files, output_dir)
                download_label = "📥 生成物を一括ダウンロード (ZIP)" if st.session_state.generation_complete else "📥 生成済み素材を一括ダウンロード (ZIP)"
                st.download_button(
                    label=f"{download_label} ({total_mb:.0f}MB)",
                    data=zip_buffer,
                    file_name=f"video_output_{output_dir.name}.zip",
                    mime="application/zip",
                    key="dl_all",
                )
            else:
                st.info(f"📦 合計 {total_mb:.0f}MB のため、一括ダウンロードは無効です。カテゴリ別にダウンロードしてください。")

            # 個別ファイル一覧
            with st.expander("📁 生成ファイル一覧"):
                for file_path in sorted(output_dir.rglob("*")):
                    if file_path.is_file():
                        st.text(f"  {file_path.relative_to(output_dir)}")
        else:
            st.info("📥 生成が完了すると、ここにダウンロードリンクが表示されます。")
    else:
        st.info("📥 生成が完了すると、ここにダウンロードリンクが表示されます。")


def _clear_step_files(directory: Path, patterns: list[str], exclude_subdir: str | None = None) -> None:
    """再生成時に既存ファイルを削除"""
    if not directory.exists():
        return
    for pattern in patterns:
        for f in directory.glob(pattern):
            if f.is_file():
                if exclude_subdir and exclude_subdir in str(f.relative_to(directory).parts):
                    continue
                f.unlink(missing_ok=True)


def detect_step_status(output_dir: Path) -> dict[str, bool]:
    """出力ディレクトリをスキャンして各ステップの完了状態を判定"""
    status = {
        "audio": False,
        "bgm": False,
        "bg_video": False,
        "images": False,
        "timeline": False,
    }
    if not output_dir or not output_dir.exists():
        return status

    audio_dir = output_dir / "audio"
    if audio_dir.exists() and list(audio_dir.glob("*.wav")):
        status["audio"] = True

    bgm_dir = output_dir / "bgm"
    if bgm_dir.exists() and (list(bgm_dir.glob("*.mp3")) or list(bgm_dir.glob("*.wav"))):
        status["bgm"] = True

    video_bg_dir = output_dir / "videos" / "backgrounds"
    if video_bg_dir.exists() and list(video_bg_dir.glob("*.mp4")):
        status["bg_video"] = True

    image_dir = output_dir / "images"
    if image_dir.exists() and (list(image_dir.glob("*.png")) or list(image_dir.glob("*.jpg"))):
        status["images"] = True

    if (output_dir / "timeline.csv").exists():
        status["timeline"] = True
    video_dir = output_dir / "videos"
    if video_dir.exists() and list(video_dir.glob("*.mp4")):
        status["timeline"] = True

    return status


def _get_or_create_history_entry(output_dir: Path) -> dict:
    """履歴エントリを取得または作成する"""
    if st.session_state.resume_mode["enabled"] and st.session_state.resume_mode["entry"]:
        entry = st.session_state.resume_mode["entry"]
        entry["status"] = "in_progress"
    else:
        entry = create_history_entry(str(output_dir))

    st.session_state.current_history_id = entry["id"]
    add_history_entry(entry)
    return entry


def run_step_audio(script, output_dir: Path, history_entry: dict | None = None) -> dict:
    """ステップ1: 音声生成"""
    import gc

    audio_mode = st.session_state.get("audio_mode", "batch")
    audio_dir = output_dir / "audio"
    audio_dir.mkdir(exist_ok=True)

    if not st.session_state.reuse_mode["enabled"]:
        st.session_state.audio_files = {}

    # 再利用モード
    if st.session_state.reuse_mode["enabled"] and st.session_state.reuse_mode["audio_files"]:
        reused_audio = st.session_state.reuse_mode["audio_files"]
        copied_audio = {}
        for key, src_path in reused_audio.items():
            src_file = Path(src_path)
            if src_file.exists():
                dst_file = audio_dir / src_file.name
                if src_file.resolve() != dst_file.resolve():
                    shutil.copy2(src_file, dst_file)
                copied_audio[key] = str(dst_file)
        st.session_state.audio_files = copied_audio
        st.success(f"♻️ 既存の音声ファイルをコピー: {len(copied_audio)}件")

        # 一括モードの場合は再利用ファイルを結合
        import wave
        if audio_mode == "batch" and "full" not in copied_audio:
            sorted_files = []
            for key, path in copied_audio.items():
                if key != "full" and Path(path).exists():
                    try:
                        num = int(key) if isinstance(key, str) else key
                        sorted_files.append((num, path))
                    except (ValueError, TypeError):
                        pass
            sorted_files.sort(key=lambda x: x[0])
            if sorted_files:
                audio_segments = []
                for _num, wav_path in sorted_files:
                    with wave.open(wav_path, "rb") as wf:
                        audio_segments.append(wf.readframes(wf.getnframes()))
                full_audio_path = audio_dir / "full_audio.wav"
                with wave.open(str(full_audio_path), "wb") as wf:
                    wf.setnchannels(1)
                    wf.setsampwidth(2)
                    wf.setframerate(24000)
                    for segment in audio_segments:
                        wf.writeframes(segment)
                st.session_state.audio_files["full"] = str(full_audio_path)
                st.success(f"✅ 音声ファイルを1本に結合: {full_audio_path.name}")

        if history_entry:
            history_entry["progress"]["audio_generated"] = True
            history_entry["files"]["audio_files"] = dict(st.session_state.audio_files)
            add_history_entry(history_entry)
        gc.collect()
        return {"success": True, "files": dict(st.session_state.audio_files), "error": None}

    # 新規生成
    total_lines = len(script.lines) if script.lines else 0
    wait_seconds = 12
    estimated_time = total_lines * (wait_seconds + 2)
    estimated_minutes = estimated_time // 60

    if total_lines > 30:
        st.warning(f"⚠️ セリフ数: {total_lines}行 - 時間がかかります")
    st.info(f"💡 予想所要時間: 約{estimated_minutes}分（各セリフ間に{wait_seconds}秒待機）")

    progress = st.progress(0)
    status = st.empty()
    status.text("🎤 音声を生成中...")

    try:
        tts = TTSClient()

        # 既存ファイルチェック（途中再開）
        existing_audio = {}
        for wav_file in sorted(audio_dir.glob("*.wav")):
            if wav_file.name != "full_audio.wav":
                try:
                    num = int(wav_file.stem.split("_")[0])
                    existing_audio[num] = str(wav_file)
                except (ValueError, IndexError):
                    pass
        if existing_audio:
            st.info(f"♻️ 既存の音声ファイル {len(existing_audio)}件を検出（途中再開）")
            st.session_state.audio_files.update(existing_audio)

        if audio_mode == "batch":
            def update_progress(current, total, message):
                progress.progress((current + 1) / total)
                status.text(f"🎤 生成中: {current + 1}/{total} - {message}")
                if history_entry and (current + 1) % 5 == 0:
                    for wav_file in sorted(audio_dir.glob("*.wav")):
                        if wav_file.name != "full_audio.wav":
                            try:
                                num = int(wav_file.stem.split("_")[0])
                                st.session_state.audio_files[num] = str(wav_file)
                            except (ValueError, IndexError):
                                pass
                    if st.session_state.audio_files:
                        history_entry["files"]["audio_files"] = dict(st.session_state.audio_files)
                        add_history_entry(history_entry)

            output_path = audio_dir / "full_audio.wav"
            wav_path = tts.synthesize_script(
                script, output_path,
                progress_callback=update_progress,
                allow_fallback=False,
            )
            st.session_state.audio_files["full"] = str(wav_path)
        else:
            for i, line in enumerate(script.lines):
                output_path = audio_dir / f"{line.number:03d}_{line.speaker}.wav"
                wav_path = tts.synthesize(line.text, line.speaker, output_path)
                st.session_state.audio_files[line.number] = str(wav_path)
                progress.progress((i + 1) / script.total_lines)

        progress.progress(1.0)
        status.text("✅ 音声生成完了")

        if history_entry:
            history_entry["progress"]["audio_generated"] = True
            history_entry["files"]["audio_files"] = dict(st.session_state.audio_files)
            add_history_entry(history_entry)

        gc.collect()
        return {"success": True, "files": dict(st.session_state.audio_files), "error": None}

    except Exception as audio_err:
        error_str = str(audio_err)
        log_error_to_file(output_dir, "音声生成エラー", error_str, traceback.format_exc())

        error_lower = error_str.lower()
        is_rate_limit = any(x in error_lower for x in ["クォータ", "quota", "429", "rate", "limit"]) or "レート制限" in error_str
        is_partial = "一部失敗" in error_str
        if is_partial:
            st.warning(f"⚠️ {audio_err}")
        elif is_rate_limit:
            st.error("❌ 音声生成がレート制限に達しました")
        else:
            st.error(f"❌ 音声生成エラー: {audio_err}")

        # 生成済みファイルを検出して保存
        if audio_dir.exists():
            for wav_file in sorted(audio_dir.glob("*.wav")):
                if wav_file.name == "full_audio.wav":
                    st.session_state.audio_files["full"] = str(wav_file)
                else:
                    try:
                        num = int(wav_file.stem.split("_")[0])
                        st.session_state.audio_files[num] = str(wav_file)
                    except (ValueError, IndexError):
                        pass

        generated_count = len(st.session_state.audio_files)
        if generated_count > 0:
            st.info(f"💾 生成済み音声: {generated_count}件を保存しました。")
            if history_entry:
                history_entry["progress"]["audio_generated"] = True
                history_entry["files"]["audio_files"] = dict(st.session_state.audio_files)
                history_entry["error"] = history_entry.get("error", "") + f"音声生成エラー({generated_count}件生成済み); "
                add_history_entry(history_entry)

        gc.collect()
        return {"success": False, "files": dict(st.session_state.audio_files), "error": error_str}


def run_step_bgm(script, prompts, output_dir: Path, history_entry: dict | None = None) -> dict:
    """ステップ2: BGM生成"""
    import gc

    bgm_dir = output_dir / "bgm"
    bgm_dir.mkdir(exist_ok=True)
    bgm_path = None

    try:
        # 再利用モード
        if st.session_state.reuse_mode["enabled"] and st.session_state.reuse_mode["bgm"]:
            src_bgm = Path(st.session_state.reuse_mode["bgm"])
            if src_bgm.exists():
                dst_bgm = bgm_dir / src_bgm.name
                if src_bgm.resolve() != dst_bgm.resolve():
                    shutil.copy2(src_bgm, dst_bgm)
                bgm_path = dst_bgm
                st.success(f"♻️ 既存のBGMファイルをコピー: {bgm_path.name}")

        if bgm_path is None:
            status = st.empty()
            status.text("🎵 BGMを生成中...")

            total_duration = 60
            if st.session_state.audio_files:
                try:
                    if "full" in st.session_state.audio_files:
                        audio_path = st.session_state.audio_files["full"]
                    else:
                        audio_path = list(st.session_state.audio_files.values())[0]
                    from moviepy import AudioFileClip
                    clip = AudioFileClip(audio_path)
                    if "full" in st.session_state.audio_files:
                        total_duration = clip.duration
                    else:
                        total_duration = clip.duration * len(st.session_state.audio_files)
                    clip.close()
                    st.info(f"🎵 音声から長さを検出: {total_duration:.1f}秒")
                except Exception:
                    pass

            if total_duration == 60 and prompts and prompts.prompts:
                last_prompt = prompts.prompts[-1]
                total_duration = time_to_seconds(last_prompt.end_time)

            bgm_path = bgm_dir / "background_music.mp3"
            bgm_client = BeatovenClient()
            bgm_client.generate(int(total_duration), bgm_path)
            if not bgm_path.exists():
                st.warning("⚠️ BGMファイルが作成されませんでした（スキップ）")
                bgm_path = None
            else:
                st.success(f"✅ BGM生成完了: {bgm_path.name}")

    except Exception as bgm_err:
        log_error_to_file(output_dir, "BGM生成エラー", str(bgm_err), traceback.format_exc())
        st.warning(f"⚠️ BGM生成に失敗（スキップ）: {bgm_err}")
        bgm_path = None

    if history_entry:
        history_entry["progress"]["bgm_generated"] = True
        history_entry["files"]["bgm"] = str(bgm_path) if bgm_path else None
        add_history_entry(history_entry)

    gc.collect()
    return {"success": bgm_path is not None, "files": {"bgm": str(bgm_path) if bgm_path else None}, "error": None}


def run_step_bg_video(script, prompts, output_dir: Path, history_entry: dict | None = None) -> dict:
    """ステップ3: 背景動画ダウンロード"""
    import gc
    import re as _re

    video_dir = output_dir / "videos" / "backgrounds"
    video_dir.mkdir(parents=True, exist_ok=True)
    background_videos = {}

    progress = st.progress(0)
    status = st.empty()

    try:
        # 再利用モード
        if st.session_state.reuse_mode["enabled"] and st.session_state.reuse_mode.get("videos"):
            status.text("♻️ 既存の動画をコピー中...")
            for num, src_path in st.session_state.reuse_mode["videos"].items():
                src_file = Path(src_path)
                if src_file.exists():
                    dst_file = video_dir / src_file.name
                    if src_file.resolve() != dst_file.resolve():
                        shutil.copy2(src_file, dst_file)
                    background_videos[num] = str(dst_file)
            if background_videos:
                st.success(f"♻️ 既存の動画: {len(background_videos)}本をコピーしました")

        if not background_videos:
            status.text("🎥 背景動画を検索中...")
            stock_client = StockVideoClient()

            _jp_to_en = {
                "不動産": "real estate", "投資": "investment", "マンション": "apartment building",
                "金利": "interest rate", "経済": "economy", "お金": "money finance",
                "株": "stock market", "ビジネス": "business office", "会議": "meeting",
                "グラフ": "chart graph", "上昇": "growth arrow", "下降": "decline",
                "都市": "city skyline", "建物": "building architecture", "家": "house home",
                "人": "people", "女性": "woman", "男性": "man", "笑顔": "smile happy",
                "驚": "surprised", "怒": "angry", "悲": "sad", "喜": "happy celebration",
                "炎": "fire flame", "水": "water ocean", "空": "sky clouds",
                "夜": "night city", "朝": "morning sunrise", "自然": "nature landscape",
                "テクノロジー": "technology", "コンピュータ": "computer", "データ": "data digital",
                "選挙": "election voting", "政治": "politics government", "ニュース": "news broadcast",
                "食事": "food dining", "料理": "cooking kitchen", "スーパー": "supermarket shopping",
                "工場": "factory industrial", "半導体": "semiconductor technology",
                "インフレ": "inflation economy", "価格": "price tag", "給料": "salary paycheck",
            }

            def _to_english_query(text: str) -> str:
                matches = []
                for jp, en in _jp_to_en.items():
                    if jp in text:
                        matches.append(en)
                if matches:
                    return " ".join(matches[:3])
                return "abstract background motion"

            chapter_pattern = _re.compile(r'【[\d:]+〜\s*(.+?)】')
            raw_content = st.session_state.get("prompt_raw_content", "") or st.session_state.get("script_raw_content", "")
            chapters = chapter_pattern.findall(raw_content)

            # 背景動画数: 画像枚数の約半分（最大20本）
            max_bg_videos = 20
            if prompts and prompts.prompts:
                target_count = min(max_bg_videos, max(1, len(prompts.prompts) // 2))
            else:
                target_count = max_bg_videos

            if chapters:
                # チャプター数がtarget_countより少なければ全部、多ければ間引き
                if len(chapters) <= target_count:
                    search_items = [(i + 1, _to_english_query(ch)) for i, ch in enumerate(chapters)]
                else:
                    ch_step = max(1, len(chapters) // target_count)
                    search_items = [(i + 1, _to_english_query(chapters[i])) for i in range(0, len(chapters), ch_step)][:target_count]
                st.info(f"📑 {len(chapters)}個のチャプターを検出 → {len(search_items)}本の背景動画を取得")
            elif prompts and prompts.prompts:
                step = max(1, len(prompts.prompts) // target_count)
                selected = prompts.prompts[::step][:target_count]
                search_items = [(p.number, _to_english_query(p.prompt)) for p in selected]
            elif script and script.lines:
                step = max(1, len(script.lines) // target_count)
                selected = script.lines[::step][:target_count]
                search_items = [(i + 1, _to_english_query(line.text)) for i, line in enumerate(selected)]
            else:
                search_items = [(1, "abstract background")]

            for i, (number, search_query) in enumerate(search_items):
                try:
                    status.text(f"🎥 背景動画検索中: {i + 1}/{len(search_items)}")
                    videos = stock_client.search_pexels(search_query, per_page=1)
                    if videos:
                        video_path = video_dir / f"{number:03d}_bg.mp4"
                        stock_client.download(videos[0], video_path)
                        background_videos[number] = str(video_path)
                        st.success(f"✅ 背景動画 {number} ダウンロード完了")
                    else:
                        videos = stock_client.search_pixabay(search_query, per_page=1)
                        if videos:
                            video_path = video_dir / f"{number:03d}_bg.mp4"
                            stock_client.download(videos[0], video_path)
                            background_videos[number] = str(video_path)
                            st.success(f"✅ 背景動画 {number} ダウンロード完了 (Pixabay)")
                except Exception as vid_err:
                    log_error_to_file(output_dir, f"背景動画取得エラー（{number}）", str(vid_err), traceback.format_exc())
                    st.warning(f"⚠️ 背景動画取得エラー（{number}）: {vid_err}")

                progress.progress((i + 1) / len(search_items))

            if background_videos:
                st.success(f"✅ 背景動画: {len(background_videos)}件準備完了")
            else:
                st.info("ℹ️ 背景動画なしで続行します")

    except Exception as bg_err:
        log_error_to_file(output_dir, "背景動画取得エラー", str(bg_err), traceback.format_exc())
        st.warning(f"⚠️ 背景動画の取得中にエラー: {bg_err}")

    if history_entry:
        history_entry["files"]["videos"] = {str(k): v for k, v in background_videos.items()}
        add_history_entry(history_entry)

    gc.collect()
    return {"success": bool(background_videos), "files": {"videos": background_videos}, "error": None}


def run_step_images(script, prompts, output_dir: Path, history_entry: dict | None = None) -> dict:
    """ステップ4: 画像生成（3枚毎にチェックポイント+gc.collect）"""
    import gc

    generated_images = {}
    reused_count = 0
    generated_count = 0

    progress = st.progress(0)
    status = st.empty()

    try:
        # 再利用モード
        if st.session_state.reuse_mode["enabled"] and st.session_state.reuse_mode["images"]:
            status.text("♻️ 既存の画像をコピー中...")
            image_dir = output_dir / "images"
            image_dir.mkdir(exist_ok=True)
            for num, src_path in st.session_state.reuse_mode["images"].items():
                src_file = Path(src_path)
                if src_file.exists():
                    dst_file = image_dir / src_file.name
                    if src_file.resolve() != dst_file.resolve():
                        shutil.copy2(src_file, dst_file)
                    generated_images[num] = str(dst_file)
                else:
                    generated_images[num] = src_path
            reused_count = len(generated_images)
            st.success(f"♻️ 既存の画像: {reused_count}枚をコピーしました")

        # 画像プロンプト自動生成
        user_specified = st.session_state.get("user_specified_num_images", 0)
        should_regenerate = prompts.total_images == 0

        if should_regenerate:
            if user_specified > 0:
                calculated_images = user_specified
            else:
                raw_content = st.session_state.get("script_raw_content", "")
                if raw_content:
                    detected_items = count_script_items_from_content(raw_content)
                else:
                    detected_items = count_script_items(script)
                if detected_items == 0:
                    detected_items = script.total_lines
                if detected_items <= 0:
                    detected_items = max(1, len(script.lines) if script.lines else 1)
                calculated_images = max(1, min(detected_items, 100))
            st.info(f"🎨 {calculated_images}件の画像プロンプトを自動生成中...")
            try:
                auto_prompts = generate_image_prompts_from_script(script, calculated_images)
                prompts = auto_prompts
                st.session_state.prompts = auto_prompts
                st.success(f"✅ {prompts.total_images}件の画像プロンプトを自動生成しました")
            except Exception as auto_err:
                st.warning(f"⚠️ 画像プロンプト自動生成エラー: {auto_err}")

        # 画像生成
        if prompts.total_images > 0:
            missing_prompts = [p for p in prompts.prompts if p.number not in generated_images]

            if missing_prompts:
                st.info(f"🖼️ 不足している画像: {len(missing_prompts)}枚を新規生成します...")
                image_gen = ImageGenerator()
                image_dir = output_dir / "images"
                image_dir.mkdir(exist_ok=True)
                stock_client = StockVideoClient()

                image_errors = []
                for i, p in enumerate(missing_prompts):
                    try:
                        status.text(f"🖼️ 画像生成中: {i + 1}/{len(missing_prompts)} - {p.prompt[:30]}...")
                        output_path = image_dir / f"{p.number:03d}_scene.png"
                        image_gen.generate(p.prompt, output_path)
                        generated_images[p.number] = str(output_path)
                        generated_count += 1
                    except Exception as img_err:
                        log_error_to_file(output_dir, f"画像生成エラー（画像 {p.number}）", str(img_err), traceback.format_exc())
                        image_errors.append(p.number)
                        try:
                            stock_path = image_dir / f"{p.number:03d}_stock.jpg"
                            keywords = p.prompt.split()[:3]
                            search_query = " ".join(keywords) if keywords else "background"
                            stock_client.download_image(search_query, stock_path)
                            generated_images[p.number] = str(stock_path)
                            generated_count += 1
                        except Exception as stock_err:
                            log_error_to_file(output_dir, f"ストック画像取得エラー（画像 {p.number}）", str(stock_err), traceback.format_exc())

                    progress.progress((i + 1) / len(missing_prompts))

                    # 3枚毎にチェックポイント + gc.collect
                    if (i + 1) % 3 == 0:
                        if history_entry:
                            history_entry["files"]["images"] = {str(k): v for k, v in generated_images.items()}
                            add_history_entry(history_entry)
                        gc.collect()

                if image_errors:
                    st.warning(f"⚠️ {len(image_errors)}枚の画像でAI生成エラー（ストック代替済み）: {image_errors}")
            else:
                st.success(f"♻️ 全ての画像が既存のものを再利用（{reused_count}枚）")

            if generated_images:
                st.success(f"✅ 画像準備完了: 再利用 {reused_count}枚 + 新規 {generated_count}枚 = 合計 {len(generated_images)}枚")
            else:
                st.error("❌ 画像を生成できませんでした")
        else:
            st.error("❌ 画像プロンプトがないため、画像生成をスキップしました")

    except Exception as img_err:
        log_error_to_file(output_dir, "画像生成エラー", str(img_err), traceback.format_exc())
        st.error(f"❌ 画像生成エラー: {img_err}")

    if history_entry:
        history_entry["progress"]["images_generated"] = True
        history_entry["files"]["images"] = {str(k): v for k, v in generated_images.items()}
        add_history_entry(history_entry)

    gc.collect()
    return {"success": bool(generated_images), "files": {"images": generated_images}, "error": None}


def run_step_timeline(script, prompts, mode: str, output_formats: list, output_dir: Path, generated_images: dict, background_videos: dict, bgm_path, history_entry: dict | None = None) -> dict:
    """ステップ5: タイムライン生成/動画合成"""
    import gc

    status = st.empty()

    def get_audio_duration_safe(audio_path: str) -> float:
        try:
            from moviepy import AudioFileClip
            clip = AudioFileClip(audio_path)
            duration = clip.duration
            clip.close()
            return duration if duration else 5.0
        except Exception:
            try:
                file_size = os.path.getsize(audio_path)
                return max(1.0, file_size / 48000)
            except Exception:
                return 5.0

    if "Filmora" in mode:
        status.text("📋 タイムラインを生成中...")
        timeline = Timeline()

        audio_total_duration = 0.0
        if st.session_state.audio_files:
            if "full" in st.session_state.audio_files:
                audio_path = st.session_state.audio_files["full"]
                duration = get_audio_duration_safe(audio_path)
                timeline.add_entry(TimelineEntry(
                    start_time=0.0, end_time=duration,
                    media_type="audio", file_path=audio_path, speaker="all",
                ))
                audio_total_duration = duration
            else:
                current_time = 0.0
                for line in script.lines:
                    if line.number in st.session_state.audio_files:
                        audio_path = st.session_state.audio_files[line.number]
                        duration = get_audio_duration_safe(audio_path)
                        timeline.add_entry(TimelineEntry(
                            start_time=current_time, end_time=current_time + duration,
                            media_type="audio", file_path=audio_path, speaker=line.speaker,
                        ))
                        current_time += duration
                audio_total_duration = current_time

        if prompts and prompts.prompts:
            last_prompt = prompts.prompts[-1]
            prompt_total_duration = time_to_seconds(last_prompt.end_time)
        else:
            prompt_total_duration = len(generated_images) * 5.0 if generated_images else 10.0

        if audio_total_duration > 0 and prompt_total_duration > 0:
            time_scale = audio_total_duration / prompt_total_duration
        else:
            time_scale = 1.0

        if prompts:
            for p in prompts.prompts:
                if p.number in generated_images:
                    scaled_start = time_to_seconds(p.start_time) * time_scale
                    scaled_end = time_to_seconds(p.end_time) * time_scale
                    timeline.add_entry(TimelineEntry(
                        start_time=scaled_start, end_time=scaled_end,
                        media_type="image", file_path=generated_images[p.number],
                    ))

        if bgm_path and Path(bgm_path).exists():
            timeline.add_entry(TimelineEntry(
                start_time=0, end_time=timeline.total_duration,
                media_type="bgm", file_path=str(bgm_path),
            ))

        timeline.to_csv(output_dir / "timeline.csv")
        status.text("✅ タイムライン生成完了")

    else:
        # 自動モード: 動画合成
        status.text("🎬 動画を合成中...")
        editor = VideoEditor()
        timeline = Timeline()

        if not st.session_state.audio_files:
            st.warning("⚠️ 音声がないため、動画合成をスキップします。")
            if history_entry:
                history_entry["status"] = "completed"
                history_entry["error"] = history_entry.get("error", "") + "音声なしのため動画合成スキップ"
                add_history_entry(history_entry)
            gc.collect()
            return {"success": True, "files": {}, "error": "音声なしのため動画合成スキップ"}

        if "full" in st.session_state.audio_files:
            audio_path = st.session_state.audio_files["full"]
            duration = get_audio_duration_safe(audio_path)
            timeline.add_entry(TimelineEntry(
                start_time=0.0, end_time=duration,
                media_type="audio", file_path=audio_path, speaker="all",
            ))
        else:
            current_time = 0.0
            for line in script.lines:
                if line.number in st.session_state.audio_files:
                    audio_path = st.session_state.audio_files[line.number]
                    duration = get_audio_duration_safe(audio_path)
                    timeline.add_entry(TimelineEntry(
                        start_time=current_time, end_time=current_time + duration,
                        media_type="audio", file_path=audio_path, speaker=line.speaker,
                    ))
                    current_time += duration

        audio_total_duration = timeline.total_duration
        if prompts and prompts.prompts:
            last_prompt = prompts.prompts[-1]
            prompt_total_duration = time_to_seconds(last_prompt.end_time)
        else:
            prompt_total_duration = audio_total_duration

        if prompt_total_duration > 0:
            time_scale = audio_total_duration / prompt_total_duration
        else:
            time_scale = 1.0

        st.info(f"📊 タイミング調整: 音声 {audio_total_duration:.1f}秒 / プロンプト {prompt_total_duration:.1f}秒 = スケール {time_scale:.2f}x")

        if prompts:
            for p in prompts.prompts:
                if p.number in generated_images:
                    scaled_start = time_to_seconds(p.start_time) * time_scale
                    scaled_end = time_to_seconds(p.end_time) * time_scale
                    if p.number in background_videos:
                        timeline.add_entry(TimelineEntry(
                            start_time=scaled_start, end_time=scaled_end,
                            media_type="video", file_path=background_videos[p.number],
                        ))
                    timeline.add_entry(TimelineEntry(
                        start_time=scaled_start, end_time=scaled_end,
                        media_type="image", file_path=generated_images[p.number],
                    ))

        if not generated_images:
            st.error("❌ 画像が生成されていないため、動画を作成できません。")
            gc.collect()
            return {"success": False, "files": {}, "error": "画像なし"}

        if not output_formats:
            st.error("❌ 出力形式が選択されていません。")
            gc.collect()
            return {"success": False, "files": {}, "error": "出力形式未選択"}

        video_dir = output_dir / "videos"
        video_dir.mkdir(exist_ok=True)

        for i, fmt in enumerate(output_formats):
            status.text(f"🎬 動画を合成中... ({i+1}/{len(output_formats)}: {fmt})")
            output_path = video_dir / f"{fmt}.mp4"
            try:
                editor.create_video(
                    timeline=timeline, output_path=output_path,
                    format_name=fmt, bgm_path=bgm_path,
                )
                st.success(f"✅ {fmt}.mp4 を生成しました")
            except Exception as video_err:
                log_error_to_file(output_dir, f"動画生成エラー（{fmt}）", str(video_err), traceback.format_exc())
                st.error(f"❌ {fmt} 動画生成エラー: {video_err}")

    if history_entry:
        history_entry["progress"]["video_generated"] = True
        history_entry["status"] = "completed"
        add_history_entry(history_entry)

    gc.collect()
    return {"success": True, "files": {}, "error": None}


def run_generation(script, prompts, mode: str, output_formats: list, generate_audio: bool = True, generate_images: bool = True, generate_bgm: bool = False, generate_bg_video: bool = False) -> None:
    """生成処理を実行（全ステップを順番に実行するラッパー）

    Args:
        generate_audio: 音声を生成するかどうか
        generate_images: 画像を生成するかどうか
        generate_bgm: BGMを生成するかどうか
        generate_bg_video: 背景動画を取得するかどうか
    """
    # デバッグ: 選択されたモードを表示
    materials_info = []
    if generate_audio:
        materials_info.append("音声")
    if generate_images:
        materials_info.append("画像")
    if generate_bgm:
        materials_info.append("BGM")
    if generate_bg_video:
        materials_info.append("背景動画")

    if "Filmora" in mode:
        st.info(f"📂 **Filmoraモード**で実行中（素材のみ出力）- 生成対象: {', '.join(materials_info) if materials_info else 'なし'}")
    else:
        st.info(f"🎬 **自動モード**で実行中（動画を生成します）: {output_formats}")

    # 出力ディレクトリを最初に作成
    output_dir = st.session_state.output_dir or get_output_dir()
    st.session_state.output_dir = output_dir

    # 履歴エントリを作成
    history_entry = None
    try:
        history_entry = _get_or_create_history_entry(output_dir)
        history_entry["settings"]["output_mode"] = mode
        history_entry["settings"]["output_formats"] = output_formats
        add_history_entry(history_entry)
    except Exception as init_err:
        st.warning(f"⚠️ 履歴初期化エラー: {init_err}")

    try:
        # 早期バリデーション
        if not script or not script.lines or len(script.lines) == 0:
            st.error("❌ 台本が空です。セリフが含まれるファイルをアップロードしてください。")
            if history_entry:
                history_entry["status"] = "interrupted"
                history_entry["error"] = "台本が空"
                add_history_entry(history_entry)
            return

        # 台本とプロンプトを保存
        if history_entry:
            history_entry["progress"]["script_parsed"] = True
            script_file = save_script_to_output(script, output_dir)
            if script_file:
                history_entry["files"]["script_file"] = str(script_file)
            if prompts:
                prompts_file = save_prompts_to_output(prompts, output_dir)
                if prompts_file:
                    history_entry["files"]["prompts_file"] = str(prompts_file)
            add_history_entry(history_entry)

        overall_progress = st.progress(0)

        # STEP 1: 音声生成
        if generate_audio:
            run_step_audio(script, output_dir, history_entry)
        else:
            st.info("⏭️ 音声生成をスキップしました")
        overall_progress.progress(0.25)

        # STEP 2: BGM生成
        bgm_result = {"files": {"bgm": None}}
        if generate_bgm:
            bgm_result = run_step_bgm(script, prompts, output_dir, history_entry)
        else:
            st.info("⏭️ BGM生成をスキップしました")
            # 再利用モードの場合は既存のBGMをコピー
            if st.session_state.reuse_mode["enabled"] and st.session_state.reuse_mode["bgm"]:
                bgm_dir = output_dir / "bgm"
                bgm_dir.mkdir(exist_ok=True)
                src_bgm = Path(st.session_state.reuse_mode["bgm"])
                if src_bgm.exists():
                    dst_bgm = bgm_dir / src_bgm.name
                    if src_bgm.resolve() != dst_bgm.resolve():
                        shutil.copy2(src_bgm, dst_bgm)
                    bgm_result["files"]["bgm"] = str(dst_bgm)
                    st.success(f"♻️ 既存のBGMファイルをコピー: {dst_bgm.name}")
        overall_progress.progress(0.35)

        # STEP 3: 背景動画
        bg_result = {"files": {"videos": {}}}
        if generate_bg_video:
            bg_result = run_step_bg_video(script, prompts, output_dir, history_entry)
        else:
            st.info("⏭️ 背景動画の取得をスキップしました")
        overall_progress.progress(0.5)

        # STEP 4: 画像生成
        img_result = {"files": {"images": {}}}
        if generate_images:
            img_result = run_step_images(script, prompts, output_dir, history_entry)
        else:
            st.info("⏭️ 画像生成をスキップしました")
        overall_progress.progress(0.75)

        # STEP 5: タイムライン/動画合成
        bgm_path_str = bgm_result["files"].get("bgm")
        bgm_path = Path(bgm_path_str) if bgm_path_str else None
        generated_images = img_result["files"].get("images", {})
        background_videos = bg_result["files"].get("videos", {})

        run_step_timeline(
            script, prompts, mode, output_formats, output_dir,
            generated_images, background_videos, bgm_path, history_entry,
        )
        overall_progress.progress(1.0)

        # 再開モードをリセット
        st.session_state.resume_mode = {"enabled": False, "entry": None}
        st.session_state.current_history_id = None
        st.session_state.generation_complete = True
        st.rerun()

    except BaseException as e:
        error_msg = str(e)
        error_trace = traceback.format_exc()

        # Streamlitの内部リラン例外はそのまま再送出
        exception_type = type(e).__name__
        if "Rerun" in exception_type or "StopException" in exception_type:
            if history_entry and history_entry.get("status") == "in_progress":
                history_entry["status"] = "interrupted"
                history_entry["error"] = "Streamlitセッション中断"
                try:
                    add_history_entry(history_entry)
                except Exception:
                    pass
            raise

        if output_dir:
            log_error_to_file(output_dir, "生成処理エラー", error_msg, error_trace)

        st.error(f"❌ 生成エラー: {error_msg}")
        st.code(error_trace)

        if history_entry:
            history_entry["status"] = "interrupted"
            history_entry["error"] = error_msg
            history_entry["error_trace"] = error_trace[:500]
            add_history_entry(history_entry)
            st.warning("⚠️ 生成が中断されました。「📜 生成履歴」から再開できます。")

    finally:
        if history_entry and history_entry.get("status") == "in_progress":
            history_entry["status"] = "interrupted"
            history_entry["error"] = "予期せぬ中断"
            if output_dir:
                log_error_to_file(output_dir, "予期せぬ中断", "finallyブロックで検出", traceback.format_exc())
            try:
                add_history_entry(history_entry)
            except Exception:
                pass


def settings_page() -> None:
    """P-002: 設定ページ"""
    st.title("⚙️ 設定")

    settings = load_settings()

    # タブで設定カテゴリを分割
    tab1, tab2, tab3, tab4 = st.tabs(["🎤 話者設定", "🔑 APIキー設定", "📁 デフォルト設定", "👤 解説者イラスト"])

    with tab1:
        st.header("話者設定")

        speakers = settings.get("speakers", {})

        st.info("💡 台本で `speaker1:` と `speaker2:` で使い分けます。表示名はキャラクター名として自由に設定できます。")

        col1, col2 = st.columns(2)

        sp1 = speakers.get("speaker1", {})
        sp2 = speakers.get("speaker2", {})

        with col1:
            sp1_current_name = sp1.get("display_name", "ナレーター1")
            st.subheader(f"🔵 speaker1 → {sp1_current_name}")
            sp1_name = st.text_input("キャラクター名", value=sp1_current_name, key="sp1_name")
            sp1_voice = st.selectbox(
                "音声",
                ["ja-JP-Neural2-B (女性)", "ja-JP-Neural2-C (男性)", "ja-JP-Neural2-D (男性)", "ja-JP-Wavenet-A (女性)"],
                index=0,
                key="sp1_voice",
            )

        with col2:
            sp2_current_name = sp2.get("display_name", "ナレーター2")
            st.subheader(f"🟠 speaker2 → {sp2_current_name}")
            sp2_name = st.text_input("キャラクター名", value=sp2_current_name, key="sp2_name")
            sp2_voice = st.selectbox(
                "音声",
                ["ja-JP-Neural2-B (女性)", "ja-JP-Neural2-C (男性)", "ja-JP-Neural2-D (男性)", "ja-JP-Wavenet-A (女性)"],
                index=1,
                key="sp2_voice",
            )

    with tab2:
        st.header("APIキー設定")
        st.warning("⚠️ APIキーは`.env`ファイルで管理することを推奨します。")

        st.markdown("""
        **必要なAPIキー:**
        1. **Google Cloud** - TTS音声生成 + Gemini画像生成
        2. **Beatoven.ai** - BGM生成
        3. **Pexels** - 動画素材取得
        4. **Pixabay** - 動画素材取得（予備）

        詳細は `.env.example` を参照してください。
        """)

        # 設定状況の確認
        st.subheader("現在の設定状況")

        api_status = {
            "GCP認証情報 (TTS)": "✅ 設定済み" if get_gcp_credentials() else "❌ 未設定",
            "GOOGLE_API_KEY": "✅ 設定済み" if get_env_var("GOOGLE_API_KEY") else "❌ 未設定",
            "BEATOVEN_API_KEY": "✅ 設定済み" if get_env_var("BEATOVEN_API_KEY") else "❌ 未設定",
            "PEXELS_API_KEY": "✅ 設定済み" if get_env_var("PEXELS_API_KEY") else "❌ 未設定",
            "PIXABAY_API_KEY": "✅ 設定済み" if get_env_var("PIXABAY_API_KEY") else "❌ 未設定",
        }

        for key, status in api_status.items():
            st.text(f"{key}: {status}")

    with tab3:
        st.header("デフォルト設定")

        defaults = settings.get("defaults", {})

        st.subheader("出力設定")
        default_format = st.multiselect(
            "デフォルト出力形式",
            ["youtube", "instagram_reel", "instagram_feed", "tiktok"],
            default=defaults.get("output_format", ["youtube"]),
        )

        st.subheader("BGM設定")
        bgm_settings = defaults.get("bgm", {})
        bgm_mood = st.selectbox(
            "デフォルトムード",
            ["neutral", "happy", "sad", "energetic", "calm"],
            index=["neutral", "happy", "sad", "energetic", "calm"].index(bgm_settings.get("mood", "neutral")),
        )
        bgm_genre = st.selectbox(
            "デフォルトジャンル",
            ["background", "corporate", "cinematic", "electronic", "acoustic"],
            index=["background", "corporate", "cinematic", "electronic", "acoustic"].index(bgm_settings.get("genre", "background")),
        )

        st.subheader("出力フォルダ")
        st.info("💡 このPCで使用するデフォルトの出力先を選択してください。「推奨」を選ぶとローカルに保存されます。")

        # ユーザーのホームディレクトリを取得
        home_dir = os.path.expanduser("~")

        # OS別のデフォルトフォルダを取得
        default_local_folder = get_default_output_folder()

        # プリセットオプション（推奨をトップに）
        preset_paths = {
            "推奨（ローカル保存）": default_local_folder,
            "ドキュメント": os.path.join(home_dir, "Documents"),
            "デスクトップ": os.path.join(home_dir, "Desktop"),
            "ダウンロード": os.path.join(home_dir, "Downloads"),
            "ホーム": home_dir,
            "相対パス (output)": "output",
            "カスタム入力": "_custom_",
        }

        # 現在の設定値からプリセットを判定
        current_folder = defaults.get("output_folder", "output")
        current_preset = "カスタム入力"
        for name, path in preset_paths.items():
            if path == current_folder:
                current_preset = name
                break

        preset_options = list(preset_paths.keys())
        selected_preset = st.selectbox(
            "出力先を選択",
            options=preset_options,
            index=preset_options.index(current_preset) if current_preset in preset_options else 0,
            key="settings_output_preset",
        )

        if selected_preset == "カスタム入力":
            output_folder = st.text_input(
                "カスタムパスを入力",
                value=current_folder if current_folder not in preset_paths.values() else "",
                key="settings_custom_output",
            )
        else:
            output_folder = preset_paths[selected_preset]
            st.text(f"📁 {output_folder}")

    with tab4:
        st.header("解説者イラスト設定")
        st.markdown("動画の左下・右下に表示する解説者キャラクターのイラストを設定します。")
        st.info("💡 台本の `speaker1:` `speaker2:` に対応するキャラクターのイラストを設定してください。")

        # 解説者イラストの保存ディレクトリ（永続的な場所）
        avatar_dir = get_persistent_avatar_dir()

        # 表示名を取得
        sp1_display = settings.get("speakers", {}).get("speaker1", {}).get("display_name", "未設定")
        sp2_display = settings.get("speakers", {}).get("speaker2", {}).get("display_name", "未設定")

        col1, col2 = st.columns(2)

        with col1:
            st.subheader("🔵 speaker1（左下に表示）")
            st.caption(f"キャラクター名: **{sp1_display}**")
            speaker1_settings = settings.get("speakers", {}).get("speaker1", {})
            speaker1_avatar = speaker1_settings.get("avatar_path", "")
            speaker1_base64 = speaker1_settings.get("avatar_base64", "")

            # 現在のイラストを表示
            if speaker1_avatar and Path(speaker1_avatar).exists():
                st.image(speaker1_avatar, width=150, caption=f"{sp1_display} のイラスト")
                st.caption("✅ 設定に保存済み" if speaker1_base64 else "⚠️ 未保存（再アップロード推奨）")
            elif speaker1_base64:
                # Base64から表示（ファイルが消えている場合）
                st.image(base64.b64decode(speaker1_base64), width=150, caption=f"{sp1_display} のイラスト（復元済み）")
                st.caption("✅ 設定から復元")
            else:
                st.info("イラスト未設定")

            # アップロード
            sp1_upload = st.file_uploader(
                f"{sp1_display} のイラストをアップロード（PNG推奨）",
                type=["png", "jpg", "jpeg", "webp"],
                key="sp1_avatar_upload",
            )
            if sp1_upload:
                ext = sp1_upload.name.split('.')[-1].lower()
                image_data = sp1_upload.getvalue()
                sp1_avatar_path = avatar_dir / f"speaker1.{ext}"

                # ファイルを保存
                with open(sp1_avatar_path, "wb") as f:
                    f.write(image_data)

                # Base64で設定に保存（永続化）
                save_avatar_to_settings("speaker1", image_data, ext)

                st.success(f"✅ アップロード完了: {sp1_avatar_path.name}")
                st.caption(f"📁 保存先: {sp1_avatar_path}")
                st.image(str(sp1_avatar_path), width=150)
                st.rerun()  # 状態を更新

        with col2:
            st.subheader("🟠 speaker2（右下に表示）")
            st.caption(f"キャラクター名: **{sp2_display}**")
            speaker2_settings = settings.get("speakers", {}).get("speaker2", {})
            speaker2_avatar = speaker2_settings.get("avatar_path", "")
            speaker2_base64 = speaker2_settings.get("avatar_base64", "")

            # 現在のイラストを表示
            if speaker2_avatar and Path(speaker2_avatar).exists():
                st.image(speaker2_avatar, width=150, caption=f"{sp2_display} のイラスト")
                st.caption("✅ 設定に保存済み" if speaker2_base64 else "⚠️ 未保存（再アップロード推奨）")
            elif speaker2_base64:
                # Base64から表示（ファイルが消えている場合）
                st.image(base64.b64decode(speaker2_base64), width=150, caption=f"{sp2_display} のイラスト（復元済み）")
                st.caption("✅ 設定から復元")
            else:
                st.info("イラスト未設定")

            # アップロード
            sp2_upload = st.file_uploader(
                f"{sp2_display} のイラストをアップロード（PNG推奨）",
                type=["png", "jpg", "jpeg", "webp"],
                key="sp2_avatar_upload",
            )
            if sp2_upload:
                ext = sp2_upload.name.split('.')[-1].lower()
                image_data = sp2_upload.getvalue()
                sp2_avatar_path = avatar_dir / f"speaker2.{ext}"

                # ファイルを保存
                with open(sp2_avatar_path, "wb") as f:
                    f.write(image_data)

                # Base64で設定に保存（永続化）
                save_avatar_to_settings("speaker2", image_data, ext)

                st.success(f"✅ アップロード完了: {sp2_avatar_path.name}")
                st.caption(f"📁 保存先: {sp2_avatar_path}")
                st.image(str(sp2_avatar_path), width=150)
                st.rerun()  # 状態を更新

        st.divider()
        st.markdown("""
        **表示仕様:**
        - 両方のキャラクターが常に表示されます
        - 話している方がハイライト（明るく）表示されます
        - 話していない方は半透明で表示されます

        💡 キャラクター名は「話者設定」タブで変更できます。
        """)

    # 保存ボタン
    st.divider()
    if st.button("💾 設定を保存", type="primary"):
        # 設定を更新
        voice_map = {
            "ja-JP-Neural2-B (女性)": "ja-JP-Neural2-B",
            "ja-JP-Neural2-C (男性)": "ja-JP-Neural2-C",
            "ja-JP-Neural2-D (男性)": "ja-JP-Neural2-D",
            "ja-JP-Wavenet-A (女性)": "ja-JP-Wavenet-A",
        }

        if "speakers" not in settings:
            settings["speakers"] = {"speaker1": {}, "speaker2": {}}

        settings["speakers"]["speaker1"]["display_name"] = sp1_name
        settings["speakers"]["speaker1"]["voice_name"] = voice_map.get(sp1_voice, "ja-JP-Neural2-B")
        settings["speakers"]["speaker2"]["display_name"] = sp2_name
        settings["speakers"]["speaker2"]["voice_name"] = voice_map.get(sp2_voice, "ja-JP-Neural2-C")

        # アバターパスを保存
        avatar_dir = Path("assets/avatars")
        for sp_key, sp_num in [("speaker1", 1), ("speaker2", 2)]:
            for ext in ["png", "jpg", "jpeg", "webp"]:
                avatar_path = avatar_dir / f"speaker{sp_num}.{ext}"
                if avatar_path.exists():
                    settings["speakers"][sp_key]["avatar_path"] = str(avatar_path)
                    break

        if "defaults" not in settings:
            settings["defaults"] = {"bgm": {}}

        settings["defaults"]["output_format"] = default_format
        if "bgm" not in settings["defaults"]:
            settings["defaults"]["bgm"] = {}
        settings["defaults"]["bgm"]["mood"] = bgm_mood
        settings["defaults"]["bgm"]["genre"] = bgm_genre
        settings["defaults"]["output_folder"] = output_folder

        save_settings(settings)
        st.success("✅ 設定を保存しました！")


def main() -> None:
    """メイン関数"""
    # サイドバーでページ選択
    with st.sidebar:
        st.title("🎬 動画生成")
        st.divider()

        page = st.radio(
            "ページを選択",
            ["🏠 動画生成メイン", "⚙️ 設定"],
            label_visibility="collapsed",
        )

        st.divider()
        st.markdown("**バージョン:** 0.2.4")
        st.markdown("[📖 ドキュメント](docs/requirements.md)")

    # ページルーティング
    if page == "🏠 動画生成メイン":
        main_page()
    else:
        settings_page()


if __name__ == "__main__":
    main()
