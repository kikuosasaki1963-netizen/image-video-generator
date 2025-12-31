"""Word ドキュメントリーダー"""
from pathlib import Path
from docx import Document


def read_word_file(file_path: str | Path) -> str:
    """
    Word (.docx) ファイルからテキストを読み込む

    Args:
        file_path: Wordファイルのパス

    Returns:
        抽出されたテキスト内容
    """
    path = Path(file_path)

    if not path.exists():
        raise FileNotFoundError(f"ファイルが見つかりません: {path}")

    if path.suffix.lower() != ".docx":
        raise ValueError(f"非対応のファイル形式: {path.suffix}。.docx のみ対応しています。")

    doc = Document(path)
    paragraphs = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    # テーブルからもテキストを抽出
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_text.append(cell_text)
            if row_text:
                paragraphs.append("\t".join(row_text))

    return "\n".join(paragraphs)
